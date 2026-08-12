"""
documents.py
------------
Document rendering for downloadable session records (PDF + Word).

Pure formatting helpers: each takes an already-created document object (an fpdf
``FPDF`` or a python-docx ``Document``) plus plain data (a summary ``dict`` /
message rows). They touch no database, no Flask session, and no shared state —
so they are safe to unit-test in isolation and safe to call from the transcript
builders in the main app.

The expected ``summary`` shape (all keys optional):
    {
      "disclaimer":      str,
      "clinical":        str,
      "codes":           [{"code", "label", "source"}],
      "codes_rationale": str,
      "copilot_cards":   [{"type", "text", "code"}],
      "client_recap":    str,
    }
"""

import io

_CARD_LABEL = {
    "risk": "Risk", "reference": "Reference",
    "suggestion": "Suggestion", "observation": "Observation",
}

# Brand palette (RGB). Used across both PDF and DOCX for one consistent look.
_BRAND    = (14, 110, 85)     # deep teal-green — header banner
_INK      = (33, 37, 41)      # body text
_MUTED    = (120, 120, 120)   # secondary text
_RULE     = (222, 228, 225)   # hairline rules
_PANEL    = (243, 247, 245)   # light panel fill
_HEADING  = (14, 110, 85)     # section headings
_PRIVATE  = (146, 39, 15)     # rust — "private / therapist-only"
_AI_GREEN = (30, 120, 60)

# Distinct colours assigned to human participants, in order.
_PARTICIPANT_RGB = [
    (30, 80, 160), (146, 39, 15), (107, 33, 168),
    (13, 110, 110), (180, 90, 9), (26, 92, 46),
]

# Consecutive messages from one speaker are grouped under a single heading, with
# a start–end time range — so a run of short messages reads as paragraphs, not a
# heading per line. A gap longer than this starts a new block.
_GROUP_GAP_SECONDS = 5 * 60


def _group_messages(messages):
    """Group consecutive messages from the same speaker into one block, starting
    a new block when the speaker changes OR the gap since the previous message
    exceeds _GROUP_GAP_SECONDS. Each group: {user_id, display_name, is_ai,
    texts: [...], start, end}."""
    groups = []
    for m in messages:
        prev = groups[-1] if groups else None
        cont = False
        if prev and prev["user_id"] == m.user_id:
            try:
                cont = (m.timestamp - prev["end"]).total_seconds() <= _GROUP_GAP_SECONDS
            except Exception:
                cont = True   # unorderable timestamps → keep them together
        if cont:
            prev["texts"].append(m.text)
            prev["end"] = m.timestamp
        else:
            groups.append({
                "user_id": m.user_id, "display_name": m.display_name,
                "is_ai": m.user_id == "AI", "texts": [m.text],
                "start": m.timestamp, "end": m.timestamp,
            })
    return groups


def _group_time_label(g):
    """'YYYY-MM-DD HH:MM' for a single-minute block, else '… HH:MM–HH:MM'."""
    start = g["start"].strftime("%Y-%m-%d %H:%M")
    end = g["end"].strftime("%H:%M")
    return start if g["start"].strftime("%H:%M") == end else f"{start}–{end}"


def _group_speaker(session_id, g):
    if g["is_ai"]:
        return "AI Co-Pilot"
    if g["display_name"]:
        return f"{session_id}-{g['display_name']}"
    return "User"


# ===========================================================================
# PDF
# ===========================================================================

DEFAULT_RECORD_LABEL = "clinical record"


def _pdf_class(record_label: str = DEFAULT_RECORD_LABEL):
    """FPDF subclass adding a footer (confidentiality note + page number) on
    every page. Defined lazily so importing this module needs no fpdf."""
    from fpdf import FPDF

    class _TranscriptPDF(FPDF):
        def footer(self):
            self.set_y(-13)
            try:
                self.set_font("DejaVu", "", 8)
            except Exception:
                self.set_font("Helvetica", "", 8)
            self.set_text_color(*_MUTED)
            self.set_draw_color(*_RULE)
            self.line(20, self.get_y() - 1, 190, self.get_y() - 1)
            self.cell(0, 8, f"Confidential {record_label}", align="L")
            self.set_y(-13)
            self.cell(0, 8, f"Page {self.page_no()}", align="R")

    return _TranscriptPDF


def _pdf_heading(pdf, text, private=False):
    from fpdf.enums import XPos, YPos
    pdf.ln(1)
    pdf.set_font("DejaVu", "B", 12)
    pdf.set_text_color(*(_PRIVATE if private else _HEADING))
    pdf.cell(0, 7, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_draw_color(*_RULE)
    pdf.line(20, pdf.get_y() + 0.5, 190, pdf.get_y() + 0.5)
    pdf.ln(2.5)


def render_summary_pdf(pdf, summary: dict) -> None:
    """Render the therapist-only clinician summary onto the PDF (above the
    transcript). Does NOT draw the 'Transcript' heading — the caller does."""
    from fpdf.enums import XPos, YPos
    mc = {"new_x": XPos.LMARGIN, "new_y": YPos.NEXT}

    _pdf_heading(pdf, "Clinician summary — private (therapist only)", private=True)
    if summary.get("disclaimer"):
        pdf.set_font("DejaVu", "", 8.5)
        pdf.set_text_color(*_MUTED)
        pdf.multi_cell(0, 4.6, summary["disclaimer"], **mc)
        pdf.ln(2)

    def _block(title, body):
        if not body:
            return
        pdf.set_font("DejaVu", "B", 10.5)
        pdf.set_text_color(*_INK)
        pdf.multi_cell(0, 6, title, **mc)
        pdf.set_font("DejaVu", "", 10)
        pdf.set_text_color(*_INK)
        pdf.multi_cell(0, 5.2, body, **mc)
        pdf.ln(2.5)

    _block("Clinical summary",
           summary.get("clinical") or "(AI narrative unavailable — see transcript below.)")

    # ICD codes — always rendered (grounded data), even if the narrative failed.
    pdf.set_font("DejaVu", "B", 10.5)
    pdf.set_text_color(*_INK)
    pdf.multi_cell(0, 6, "ICD codes (billing reference)", **mc)
    pdf.set_font("DejaVu", "", 10)
    codes = summary.get("codes") or []
    if codes:
        for c in codes:
            label = f"{c['label']} — " if c.get("label") else ""
            pdf.multi_cell(0, 5.2, f"•  {label}{c.get('code', '')}  ({c.get('source', '')})", **mc)
    else:
        pdf.multi_cell(0, 5.2, "No ICD reference codes surfaced during this session.", **mc)
    if summary.get("codes_rationale"):
        pdf.ln(0.5)
        pdf.set_text_color(*_MUTED)
        pdf.multi_cell(0, 5, summary["codes_rationale"], **mc)
    pdf.ln(2.5)

    # Co-pilot alerts — the full saved record (incl. any not shown live).
    cards = summary.get("copilot_cards") or []
    if cards:
        pdf.set_font("DejaVu", "B", 10.5)
        pdf.set_text_color(*_INK)
        pdf.multi_cell(0, 6, "Co-pilot alerts (full record)", **mc)
        pdf.set_font("DejaVu", "", 10)
        for c in cards:
            lbl = _CARD_LABEL.get(c.get("type"), (c.get("type") or "Note").title())
            code = f"  [{c['code']}]" if c.get("code") else ""
            pdf.multi_cell(0, 5.2, f"•  [{lbl}] {c.get('text', '')}{code}", **mc)
        pdf.ln(2.5)

    _block("Client-facing draft — share only at your discretion "
           "(the client cannot see this unless you give it to them)",
           summary.get("client_recap"))


def transcript_pdf_buf(session_id, messages, mode, generated_at,
                       friendly_label=None, summary=None,
                       font_regular=None, font_bold=None,
                       record_label=DEFAULT_RECORD_LABEL) -> io.BytesIO:
    """Render a session transcript as a PDF in memory (data passed in).

    `record_label` is what this file IS for the practitioner whose session it is —
    a clinical record, session notes, or a recording. Passed in rather than
    assumed, so a coach's file is not labelled as clinical.
    """
    from fpdf.enums import XPos, YPos

    pdf = _pdf_class(record_label)()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(20, 20, 20)
    pdf.add_font("DejaVu",      fname=font_regular)
    pdf.add_font("DejaVu", "B", fname=font_bold)
    pdf.add_page()

    # ---- Header banner (page 1) --------------------------------------------
    pdf.set_fill_color(*_BRAND)
    pdf.rect(0, 0, 210, 26, style="F")
    pdf.set_xy(20, 7)
    pdf.set_font("DejaVu", "B", 17)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 8, "TogetherMindsAI", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_x(20)
    pdf.set_font("DejaVu", "", 9)
    pdf.cell(0, 5, "Session record — confidential; encrypted, never sold or used to train AI.",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_y(32)

    # ---- Metadata panel ----------------------------------------------------
    meta = [("Session ID", session_id)]
    if friendly_label:
        meta.append(("Session name", friendly_label))
    meta.append(("Mode", mode))
    meta.append(("Generated", generated_at))
    box_h = 5.8 * len(meta) + 5
    y0 = pdf.get_y()
    pdf.set_fill_color(*_PANEL)
    pdf.rect(20, y0, 170, box_h, style="F")
    pdf.set_xy(25, y0 + 2.5)
    for label, val in meta:
        pdf.set_font("DejaVu", "B", 9)
        pdf.set_text_color(*_MUTED)
        pdf.cell(34, 5.8, label, new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_font("DejaVu", "", 9.5)
        pdf.set_text_color(*_INK)
        pdf.cell(0, 5.8, str(val), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_x(25)
    pdf.set_y(y0 + box_h + 4)

    # ---- Therapist-only summary (Pro/Premium) ------------------------------
    if summary is not None:
        render_summary_pdf(pdf, summary)

    # ---- Transcript --------------------------------------------------------
    _pdf_heading(pdf, "Transcript")

    # Assign a distinct colour to each human participant + a short name for the legend.
    color_by_user, name_by_user, idx = {}, {}, 0
    for msg in messages:
        if msg.user_id != "AI" and msg.user_id not in color_by_user:
            color_by_user[msg.user_id] = _PARTICIPANT_RGB[idx % len(_PARTICIPANT_RGB)]
            name_by_user[msg.user_id] = msg.display_name or "User"
            idx += 1

    # Speaker colour legend.
    if color_by_user or messages:
        pdf.set_font("DejaVu", "", 8)
        pdf.set_text_color(*_MUTED)
        pdf.cell(pdf.get_string_width("Speakers:") + 3, 5, "Speakers:",
                 new_x=XPos.RIGHT, new_y=YPos.TOP)
        for uid, rgb in color_by_user.items():
            pdf.set_text_color(*rgb)
            nm = name_by_user.get(uid, "User")
            pdf.cell(pdf.get_string_width(nm) + 5, 5, nm, new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_text_color(*_AI_GREEN)
        pdf.cell(0, 5, "AI Co-Pilot", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

    if not messages:
        pdf.set_text_color(*_MUTED)
        pdf.set_font("DejaVu", "", 11)
        pdf.cell(0, 8, "No messages recorded for this session.",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    else:
        # One heading per speaker-block (see _group_messages), with a start–end
        # time range; each message in the block is its own paragraph beneath it.
        for g in _group_messages(messages):
            rgb = _AI_GREEN if g["is_ai"] else color_by_user.get(g["user_id"], (30, 80, 160))
            pdf.set_font("DejaVu", "B", 10)
            pdf.set_text_color(*rgb)
            pdf.cell(0, 6.5, f"{_group_speaker(session_id, g)}  ·  {_group_time_label(g)}",
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("DejaVu", "", 10.5)
            pdf.set_text_color(*_INK)
            for text in g["texts"]:
                pdf.multi_cell(0, 5.6, text)
                pdf.ln(1)
            pdf.ln(2)

    buf = io.BytesIO(pdf.output())
    buf.seek(0)
    return buf


# ===========================================================================
# DOCX
# ===========================================================================

def _docx_page_number(paragraph):
    """Insert a live PAGE field into a docx paragraph (renders in Word)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    for kind, txt in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = txt
        run._r.append(el)


def _docx_heading(doc, text, private=False):
    from docx.shared import Pt, RGBColor
    h = doc.add_heading(text, level=2)
    rgb = RGBColor(*(_PRIVATE if private else _HEADING))
    for r in h.runs:
        r.font.color.rgb = rgb


def render_summary_docx(doc, summary: dict) -> None:
    """Render the therapist-only clinician summary onto the DOCX (above the
    transcript). Does NOT add the 'Transcript' heading — the caller does."""
    from docx.shared import Pt, RGBColor

    _docx_heading(doc, "Clinician summary — private (therapist only)", private=True)
    if summary.get("disclaimer"):
        disc = doc.add_paragraph()
        run = disc.add_run(summary["disclaimer"])
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*_MUTED)

    def _block(title, body):
        if not body:
            return
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(body)

    _block("Clinical summary",
           summary.get("clinical") or "(AI narrative unavailable — see transcript below.)")

    p = doc.add_paragraph()
    p.add_run("ICD codes (billing reference)").bold = True
    codes = summary.get("codes") or []
    if codes:
        for c in codes:
            label = f"{c['label']} — " if c.get("label") else ""
            doc.add_paragraph(f"{label}{c.get('code', '')}  ({c.get('source', '')})", style="List Bullet")
    else:
        doc.add_paragraph("No ICD reference codes surfaced during this session.")
    if summary.get("codes_rationale"):
        rat = doc.add_paragraph()
        rr = rat.add_run(summary["codes_rationale"])
        rr.font.color.rgb = RGBColor(*_MUTED)

    cards = summary.get("copilot_cards") or []
    if cards:
        p = doc.add_paragraph()
        p.add_run("Co-pilot alerts (full record)").bold = True
        for c in cards:
            lbl = _CARD_LABEL.get(c.get("type"), (c.get("type") or "Note").title())
            code = f"  [{c['code']}]" if c.get("code") else ""
            doc.add_paragraph(f"[{lbl}] {c.get('text', '')}{code}", style="List Bullet")

    _block("Client-facing draft — share only at your discretion "
           "(the client cannot see this unless you give it to them)",
           summary.get("client_recap"))


def transcript_docx_buf(session_id, messages, mode, generated_at,
                        friendly_label=None, summary=None,
                        record_label=DEFAULT_RECORD_LABEL) -> io.BytesIO:
    """Render a session transcript as a DOCX in memory (data passed in)."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # ---- Title + confidentiality subtitle ----------------------------------
    title = doc.add_heading("TogetherMindsAI — Session Record", level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(*_BRAND)
    sub = doc.add_paragraph()
    sr = sub.add_run(f"Confidential {record_label} — encrypted; never sold or used to train AI.")
    sr.italic = True
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(*_MUTED)

    # ---- Metadata table ----------------------------------------------------
    meta = [("Session ID", session_id)]
    if friendly_label:
        meta.append(("Session name", friendly_label))
    meta.append(("Mode", mode))
    meta.append(("Generated", generated_at))
    table = doc.add_table(rows=0, cols=2)
    for label, val in meta:
        cells = table.add_row().cells
        lr = cells[0].paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(*_MUTED)
        vr = cells[1].paragraphs[0].add_run(str(val))
        vr.font.size = Pt(9)

    # ---- Therapist-only summary --------------------------------------------
    if summary is not None:
        render_summary_docx(doc, summary)

    # ---- Transcript --------------------------------------------------------
    _docx_heading(doc, "Transcript")

    color_by_user, idx = {}, 0
    for msg in messages:
        if msg.user_id != "AI" and msg.user_id not in color_by_user:
            color_by_user[msg.user_id] = RGBColor(*_PARTICIPANT_RGB[idx % len(_PARTICIPANT_RGB)])
            idx += 1

    if not messages:
        p = doc.add_paragraph("No messages recorded for this session.")
        p.runs[0].italic = True
    else:
        # One heading per speaker-block (see _group_messages), with a start–end
        # time range; each message in the block is its own paragraph beneath it.
        for g in _group_messages(messages):
            rgb = RGBColor(*_AI_GREEN) if g["is_ai"] else color_by_user.get(g["user_id"], RGBColor(30, 80, 160))
            p = doc.add_paragraph()
            run = p.add_run(f"{_group_speaker(session_id, g)}  ·  {_group_time_label(g)}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = rgb
            for text in g["texts"]:
                doc.add_paragraph(text)

    # ---- Footer: confidentiality + page number -----------------------------
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fr = fp.add_run("Confidential clinical record — Page ")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(*_MUTED)
    _docx_page_number(fp)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
