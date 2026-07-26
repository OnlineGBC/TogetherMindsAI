"""
tests/test_documents.py
-----------------------
Unit tests for documents.py — the pure summary renderers extracted from the app
monolith. They take a document object + a summary dict and touch nothing else,
so they can be exercised in isolation (no app, no DB, no Flask session).
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

import io

import documents

_FONT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "fonts")

_SUMMARY = {
    "disclaimer": "Private to the clinician.",
    "clinical": "Client reports low mood and poor sleep.",
    "codes": [{"code": "F32.1", "label": "Moderate depressive episode", "source": "ICD-11"}],
    "codes_rationale": "Grounded from the session's flagged themes.",
    "copilot_cards": [{"type": "risk", "text": "Sleep disruption noted.", "code": ""}],
    "client_recap": "You spoke about a hard week; we agreed to try a sleep routine.",
}


def test_render_summary_pdf_writes_content():
    from fpdf import FPDF
    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_font("DejaVu", fname=os.path.join(_FONT_DIR, "DejaVuSans.ttf"))
    pdf.add_font("DejaVu", "B", fname=os.path.join(_FONT_DIR, "DejaVuSans-Bold.ttf"))
    pdf.add_page()

    documents.render_summary_pdf(pdf, _SUMMARY)

    out = bytes(pdf.output())
    assert out.startswith(b"%PDF")       # a valid PDF was produced
    assert len(out) > 800                 # non-trivial content was written


def test_render_summary_docx_writes_sections():
    from docx import Document
    doc = Document()

    documents.render_summary_docx(doc, _SUMMARY)

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "ICD codes (billing reference)" in text
    assert "F32.1" in text
    assert "Co-pilot alerts (full record)" in text
    assert "Sleep disruption noted." in text
    # Round-trips to a real .docx file in memory.
    buf = io.BytesIO()
    doc.save(buf)
    assert buf.getbuffer().nbytes > 0


def test_render_summary_docx_handles_empty_summary():
    """Missing/empty keys must not raise — the record still renders."""
    from docx import Document
    doc = Document()
    documents.render_summary_docx(doc, {})
    text = "\n".join(p.text for p in doc.paragraphs)
    # ICD block is always emitted, even with no codes.
    assert "No ICD reference codes surfaced during this session." in text


# ---------------------------------------------------------------------------
# Transcript builders — pure, data passed in (no DB / session)
# ---------------------------------------------------------------------------

class _Msg:
    """Minimal stand-in for a ChatMessage row."""
    def __init__(self, user_id, display_name, text, ts):
        self.user_id = user_id
        self.display_name = display_name
        self.text = text
        self.timestamp = ts


def _messages():
    from datetime import datetime
    return [
        _Msg("u1", "Alex", "I felt anxious this week.", datetime(2026, 7, 1, 10, 0)),
        _Msg("AI", None, "Noted — a grounding exercise may help.", datetime(2026, 7, 1, 10, 1)),
    ]


def test_transcript_pdf_buf_renders_messages_and_summary():
    buf = documents.transcript_pdf_buf(
        "SESS-1", _messages(), "couple", "2026-07-01 10:05 UTC",
        friendly_label="Smith wk3", summary=_SUMMARY,
        font_regular=os.path.join(_FONT_DIR, "DejaVuSans.ttf"),
        font_bold=os.path.join(_FONT_DIR, "DejaVuSans-Bold.ttf"),
    )
    out = buf.read()
    assert out.startswith(b"%PDF")
    assert len(out) > 1000


def test_transcript_pdf_buf_handles_no_messages():
    buf = documents.transcript_pdf_buf(
        "SESS-2", [], "solo", "2026-07-01 10:05 UTC",
        font_regular=os.path.join(_FONT_DIR, "DejaVuSans.ttf"),
        font_bold=os.path.join(_FONT_DIR, "DejaVuSans-Bold.ttf"),
    )
    assert buf.read().startswith(b"%PDF")


def _docx_all_text(doc):
    """All text in a docx — paragraphs AND table cells (metadata is a table)."""
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_transcript_docx_buf_renders_speakers():
    from docx import Document
    buf = documents.transcript_docx_buf(
        "SESS-3", _messages(), "couple", "2026-07-01 10:05 UTC",
        friendly_label="Smith wk3", summary=None,
    )
    doc = Document(buf)
    text = _docx_all_text(doc)
    assert "Session name" in text and "Smith wk3" in text   # metadata table
    assert "SESS-3-Alex" in text          # human speaker labelled with session prefix
    assert "AI Co-Pilot" in text           # AI speaker labelled
    assert "I felt anxious this week." in text


def _M(u, d, t, ts):
    return _Msg(u, d, t, ts)


def test_group_messages_groups_and_breaks_on_gap():
    from datetime import datetime, timedelta
    base = datetime(2026, 7, 1, 20, 0)
    msgs = [
        _M("u1", "Alex", "one", base),
        _M("u1", "Alex", "two", base + timedelta(minutes=1)),     # same speaker, small gap -> grouped
        _M("u2", "Sam", "hi", base + timedelta(minutes=2)),       # speaker change -> new block
        _M("u1", "Alex", "later", base + timedelta(minutes=30)),  # same speaker, >5min gap -> new block
    ]
    groups = documents._group_messages(msgs)
    assert len(groups) == 3
    assert groups[0]["texts"] == ["one", "two"]
    assert groups[0]["end"] == base + timedelta(minutes=1)
    assert groups[1]["texts"] == ["hi"]
    assert groups[2]["texts"] == ["later"]                        # gap-break


def test_transcript_docx_groups_consecutive_same_speaker():
    from docx import Document
    from datetime import datetime, timedelta
    base = datetime(2026, 7, 1, 20, 0)
    msgs = [
        _M("u1", "Alex", "first sentence.", base),
        _M("u1", "Alex", "second sentence.", base + timedelta(minutes=1)),
        _M("u1", "Alex", "third sentence.", base + timedelta(minutes=2)),
    ]
    doc = Document(documents.transcript_docx_buf("SESS-G", msgs, "solo", "2026-07-01 20:10 UTC"))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert text.count("SESS-G-Alex") == 1          # ONE heading, not three
    assert "20:00–20:02" in text                    # start–end range
    assert "first sentence." in text and "third sentence." in text
