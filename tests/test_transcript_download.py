"""Tests for the transcript download routes (PDF + DOCX).

Regression: the routes previously returned the body via make_response(bytes),
which doesn't advertise range support — Brave's ranged download reset mid-transfer
("Check internet connection") and showed a generic filename. They now use
send_file (conditional/range capable) with a proper attachment filename.

These assert the HTTP-level contract: 200, attachment Content-Disposition with the
real transcript filename, correct content-type, and a non-empty body.
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

os.environ["TESTING"] = "1"
os.environ.setdefault("SECRET_KEY", "test-secret-key-download")
os.environ.setdefault("CORS_ALLOWED_ORIGINS", "http://localhost:5001")

import uuid
from datetime import datetime, timezone

import pytest
from sqlalchemy import create_engine
from sqlalchemy.pool import StaticPool

from TogetherMindsAI import app
from models import db, User, TherapySession, ChatMessage
from session_id import generate_session_id, filename_slug


@pytest.fixture()
def client():
    test_engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    db._app_engines[app] = {None: test_engine}
    app.config["TESTING"] = True
    with app.app_context():
        db.create_all()
        with app.test_client() as c:
            yield c
        db.session.remove()
        db.drop_all()


def _seed_session_with_message(user_id, session_id):
    db.session.add(User(id=user_id, therapy_mode="solo"))
    db.session.add(TherapySession(
        id=session_id, mode="solo", created_by=user_id,
        created_at=datetime.now(timezone.utc),
    ))
    db.session.add(ChatMessage(
        session_id=session_id, user_id=user_id, text="Hello there.",
        display_name="Solo1",
    ))
    db.session.commit()


@pytest.mark.parametrize("fmt,ctype", [
    ("pdf",  "application/pdf"),
    ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
])
def test_transcript_download_returns_attachment(client, fmt, ctype):
    user_id = str(uuid.uuid4())
    session_id = generate_session_id()
    with app.app_context():
        _seed_session_with_message(user_id, session_id)

    with client.session_transaction() as sess:
        sess["user_id"] = user_id

    rv = client.get(f"/transcript/{session_id}/{fmt}")

    assert rv.status_code == 200
    assert rv.headers["Content-Type"].startswith(ctype)
    disp = rv.headers.get("Content-Disposition", "")
    assert "attachment" in disp
    # Real filename is exposed (not a generic "name.pdf"): the session name/id
    # prefix (slugified) + kind + date. No friendly name here, so it's the id.
    prefix = filename_slug("", fallback=session_id)
    assert f"{prefix}_transcript_" in disp
    assert disp.rstrip('"').endswith(f".{fmt}")
    assert len(rv.data) > 0


def test_docx_includes_friendly_session_name(client):
    """The downloaded document shows the therapist's friendly Session name in
    addition to the Session ID, when one has been set."""
    import io
    from docx import Document

    user_id = str(uuid.uuid4())
    session_id = generate_session_id()
    with app.app_context():
        db.session.add(User(id=user_id, therapy_mode="solo"))
        db.session.add(TherapySession(
            id=session_id, mode="solo", created_by=user_id,
            created_at=datetime.now(timezone.utc), friendly_name="Team Alpha",
        ))
        db.session.add(ChatMessage(
            session_id=session_id, user_id=user_id, text="Hello there.",
            display_name="Solo1",
        ))
        db.session.commit()

    with client.session_transaction() as sess:
        sess["user_id"] = user_id

    rv = client.get(f"/transcript/{session_id}/docx")
    assert rv.status_code == 200
    doc = Document(io.BytesIO(rv.data))
    text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:                       # metadata is a borderless table
        for row in t.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    assert "Session name" in text
    assert "Team Alpha" in text
    # The Session ID is still present too.
    assert session_id in text


@pytest.mark.parametrize("fmt", ["pdf", "docx"])
def test_transcript_download_forbidden_without_access(client, fmt):
    """A user with no stake in the session can't download its transcript."""
    owner_id = str(uuid.uuid4())
    session_id = generate_session_id()
    with app.app_context():
        _seed_session_with_message(owner_id, session_id)

    with client.session_transaction() as sess:
        sess["user_id"] = str(uuid.uuid4())  # unrelated user

    rv = client.get(f"/transcript/{session_id}/{fmt}")
    assert rv.status_code == 403
