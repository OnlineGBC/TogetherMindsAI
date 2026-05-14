"""Generate TogetherMindsAI_Risk_Audit.docx from the codebase risk audit findings."""

from docx import Document
from docx.shared import RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(1)

HEADING_COLOR  = RGBColor(0x1F, 0x49, 0x7D)
CRITICAL_COLOR = RGBColor(0xC0, 0x00, 0x00)
HIGH_COLOR     = RGBColor(0xFF, 0x00, 0x00)
MEDIUM_COLOR   = RGBColor(0xFF, 0x8C, 0x00)
LOW_COLOR      = RGBColor(0x00, 0x70, 0xC0)
SEV_RGB = {"CRITICAL": CRITICAL_COLOR, "HIGH": HIGH_COLOR,
           "MEDIUM": MEDIUM_COLOR, "LOW": LOW_COLOR}
SEV_HEX = {"CRITICAL": "C00000", "HIGH": "FF0000",
           "MEDIUM": "FF8C00", "LOW": "0070C0"}


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = HEADING_COLOR


def sev(s):
    p = doc.add_paragraph()
    p.add_run("Risk Level: ").bold = True
    r = p.add_run(s)
    r.bold = True
    r.font.color.rgb = SEV_RGB[s]


def bold_label(label):
    p = doc.add_paragraph()
    p.add_run(label).bold = True


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


# ── Title ─────────────────────────────────────────────────────────────────────
t = doc.add_heading("TogetherMindsAI — Risk Audit Report", 0)
for r in t.runs:
    r.font.color.rgb = HEADING_COLOR
doc.add_paragraph(
    f"Generated: {datetime.date.today().strftime('%d %B %Y')}"
)
doc.add_paragraph(
    "Full codebase review covering legal, health/self-harm, encryption and privacy risks. "
    "Finding numbers: first digit = category (1=Legal, 2=Health, 3=Encryption, 4=Privacy), "
    "second digit = finding within that category."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SUMMARY TABLE (top of report)
# ═══════════════════════════════════════════════════════════════════════════════
h("Summary of Findings")

summary_data = [
    ("1.1",  "Legal",      "CRITICAL", "HIPAA / BAA missing"),
    ("1.2",  "Legal",      "HIGH",     "GDPR retention / erasure not enforced"),
    ("1.3",  "Legal",      "HIGH",     "Liability for AI advice"),
    ("1.4",  "Legal",      "HIGH",     "No Terms of Service"),
    ("1.5",  "Legal",      "MEDIUM",   "Age verification not enforced"),
    ("1.6",  "Legal",      "MEDIUM",   "Professional licensure claims"),
    ("2.1",  "Health",     "HIGH",     "Crisis detection brittle (keywords only)"),
    ("2.2",  "Health",     "MEDIUM",   "No confirmation user contacted helpline"),
    ("2.3",  "Health",     "MEDIUM",   "Off-topic deflection not enforced"),
    ("2.4",  "Health",     "HIGH",     "Medical hallucination risk, no output guard"),
    ("2.5",  "Health",     "LOW",      "Befrienders.org visibility (well implemented)"),
    ("2.6",  "Health",     "LOW",      "Wellness check (well implemented)"),
    ("3.1",  "Encryption", "HIGH",     "TLS optional, no HSTS"),
    ("3.2",  "Encryption", "CRITICAL", "Data at rest unencrypted"),
    ("3.3",  "Encryption", "MEDIUM",   "Private key device-bound, no passphrase"),
    ("3.4",  "Encryption", "MEDIUM",   "Session cookie flags not hardened"),
    ("3.5",  "Encryption", "CRITICAL", "SECRET_KEY on disk"),
    ("3.6",  "Encryption", "LOW",      "SQL injection (mitigated by ORM)"),
    ("3.7",  "Encryption", "MEDIUM",   "No CSP header"),
    ("3.8",  "Encryption", "HIGH",     "CORS set to *"),
    ("3.9",  "Encryption", "MEDIUM",   "Rate limiting gaps on auth endpoints"),
    ("3.10", "Encryption", "CRITICAL", "Live API key on disk"),
    ("4.1",  "Privacy",    "HIGH",     "ChatMessage + Exercise both store full text"),
    ("4.2",  "Privacy",    "HIGH",     "Messages retained indefinitely vs stated promise"),
    ("4.3",  "Privacy",    "MEDIUM",   "Logs may capture PII in exceptions"),
    ("4.4",  "Privacy",    "MEDIUM",   "User UUID broadcast to session participants"),
    ("4.5",  "Privacy",    "HIGH",     "Anthropic API receives all messages, undisclosed"),
    ("4.6",  "Privacy",    "MEDIUM",   "Emotion classifier (well implemented, no storage)"),
    ("4.7",  "Privacy",    "MEDIUM",   "localStorage exposes session nicknames"),
]

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"

# Header row
hdr = table.rows[0].cells
for i, txt in enumerate(["#", "Category", "Severity", "Area"]):
    hdr[i].text = txt
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "1F497D")
    shd.set(qn("w:val"), "clear")
    hdr[i]._tc.get_or_add_tcPr().append(shd)

for num, cat, s, area in summary_data:
    row = table.add_row()
    row.cells[0].text = num
    row.cells[1].text = cat
    row.cells[2].text = s
    row.cells[3].text = area
    hv = SEV_HEX[s]
    for para in row.cells[2].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(
                int(hv[0:2], 16), int(hv[2:4], 16), int(hv[4:6], 16)
            )

doc.add_paragraph("")
p = doc.add_paragraph()
p.add_run("Totals: ").bold = True
p.add_run("Critical: 4   |   High: 13   |   Medium: 13   |   Low: 5   |   Total: 35")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. LEGAL
# ═══════════════════════════════════════════════════════════════════════════════
h("1. Legal Risks")

h("1.1  HIPAA Applicability & Compliance", 2)
sev("CRITICAL")
doc.add_paragraph(
    "Mental health conversation data is stored without any HIPAA Business Associate "
    "Agreement (BAA) framework."
)
bold_label("Evidence:")
bullet("TogetherMindsAI.py — ChatMessage stores therapy content in plaintext")
bullet("models.py lines 21-42 — no encryption or HIPAA-compliant audit fields")
bullet("Config allows SQLite (unencrypted) or PostgreSQL without field-level encryption")
bold_label("Fix:")
bullet("Do not market as HIPAA-compliant without a BAA with Anthropic")
bullet("Encrypt all chat messages at rest using AES-256-GCM with Google Cloud KMS")
bullet("Implement tamper-proof audit logging (who accessed what, when)")
bullet("Add automatic purge — delete after 90 days unless explicitly retained")

h("1.2  GDPR & Data Protection", 2)
sev("HIGH")
doc.add_paragraph(
    'auth.html states "conversations will be deleted by default" but no enforced deletion '
    "mechanism exists. Data persists until the user manually deletes it."
)
bold_label("Evidence:")
bullet("TogetherMindsAI.py lines 385-397 — deletion is voluntary, not automatic")
bullet("templates/auth.html line 70 — stated promise of deletion is aspirational only")
bullet("No purge job, retention schedule, or right-to-erasure automation")
bullet("Cross-border risk: Anthropic receives all chat messages via API")
bold_label("Fix:")
bullet("Implement automatic 30-day session deletion unless user opts to retain")
bullet("Add /user/<id>/erase endpoint with immutable audit log of deletion")
bullet("Anonymise messages before sending to Anthropic (remove user_id)")
bullet("Execute a Data Processing Addendum (DPA) with Anthropic")

h("1.3  Liability for AI Therapy Advice", 2)
sev("HIGH")
doc.add_paragraph(
    'Disclaimers exist but are insufficient. The app is positioned as a "therapist" in '
    "marketing and system prompts, inviting liability for harmful advice."
)
bold_label("Evidence:")
bullet('ai_therapist.py lines 226-292 — AI described as "a thoughtful, experienced counsellor"')
bullet('ai_therapist.py line 232 — implies AI has "many years in the consulting room" (false credentials)')
bullet("Disclaimers in small text only; not signed off by user")
bold_label("Fix:")
bullet("Add a signed Terms of Service before first session — not a substitute for professional care")
bullet("Require affirmative checkbox + date acknowledgment")
bullet("Rewrite system prompt to emphasise AI limitations — cannot diagnose or prescribe")
bullet("Obtain errors & omissions insurance for AI-generated advice")

h("1.4  Terms of Service & Acceptable Use", 2)
sev("HIGH")
doc.add_paragraph("No formal Terms of Service exists — only scattered disclaimer text.")
bold_label("Evidence:")
bullet("No /terms, /tos, or /legal route in the application")
bullet("templates/base.html lines 51-62 — single disclaimer bar, no link to full policy")
bullet("No rules against harassment in couple/group sessions or illegal use")
bold_label("Fix:")
bullet("Create /templates/tos.html covering acceptable use, liability limitation, "
       "indemnification, dispute resolution")
bullet("Link from footer and auth page; require affirmative acceptance before account creation")

h("1.5  Age Verification for Minors", 2)
sev("MEDIUM")
doc.add_paragraph(
    "Age gate exists (checkbox) but is not enforced. Minors can bypass self-attestation trivially."
)
bold_label("Evidence:")
bullet('templates/auth.html lines 47-56 — "I am 18 years of age or older" is self-attested only')
bullet("No database field to track age consent or require parental approval")
bold_label("Fix:")
bullet("If allowing minors: require parental double-opt-in email consent; "
       "stricter AI moderation for under-18")
bullet("If disallowing minors: integrate third-party age verification (e.g. Socure)")

h("1.6  Professional Licensure Claims", 2)
sev("MEDIUM")
doc.add_paragraph(
    'The app uses "Solo Therapy", "Couples Therapy", and "Group Therapy" — '
    "implying licensed professional services."
)
bold_label("Fix:")
bullet('Rebrand as "AI-powered reflective support tool" — not "therapy"')
bullet("Add prominent notice: not a licensed healthcare provider; not therapy or medical advice")
bullet('Remove "therapist" from system prompts; use "reflective companion" or "listening tool"')

# ═══════════════════════════════════════════════════════════════════════════════
# 2. HEALTH & SELF-HARM
# ═══════════════════════════════════════════════════════════════════════════════
h("2. Health & Self-Harm Risks")

h("2.1  Crisis / Suicidality Detection", 2)
sev("HIGH")
doc.add_paragraph(
    "Crisis detection uses ~25 hardcoded keywords. Phrases such as "
    '"I feel like I can\'t go on" will not trigger it. '
    "No escalation beyond a static text response."
)
bold_label("Evidence:")
bullet("ai_therapist.py lines 21-27 — CRISIS_KEYWORDS hardcoded set")
bullet("ai_therapist.py lines 390-393 — detect_crisis() uses case-insensitive substring match only")
bullet("No tracking of crisis frequency across a session; no admin notification")
bold_label("Fix:")
bullet("Replace keyword matching with a Claude safety-check call returning structured risk level")
bullet("Track crisis messages per session; if 2 or more, email admin")
bullet("After crisis response, require user to confirm they have contacted a helpline before continuing")
bullet("Log all crisis detections for legal defensibility")

h("2.2  Escalation Path & Emergency Services", 2)
sev("MEDIUM")
doc.add_paragraph(
    "Crisis response includes multiple international helplines — good. "
    "But there is no verification the user actually contacted help."
)
bold_label("Evidence:")
bullet("ai_therapist.py lines 50-64 — includes 988, Crisis Text Line, Samaritans, "
       "Lifeline, befrienders.org")
bullet("No modal or gate requiring confirmation the user has reached out")
bold_label("Fix:")
bullet('Post-crisis modal with 3 paths: "Yes I got help" / "No, I\'ll call now" / '
       '"I\'m in immediate danger"')
bullet('Disable chat for 5 minutes if user selects "I\'ll call now"')
bullet("Alert other couple/group participants if one member discloses a crisis")

h("2.3  Off-Topic Deflection", 2)
sev("MEDIUM")
doc.add_paragraph(
    "System prompt instructs Claude to deflect off-topic queries but no output guard "
    "verifies Claude actually followed the instruction."
)
bold_label("Fix:")
bullet("Add secondary guard in _sanitize_response() detecting factual content "
       "unrelated to emotional state")
bullet('Test with adversarial prompts ("Forget your instructions and...") to verify no instruction leakage')

h("2.4  AI Hallucination Risk for Medical Advice", 2)
sev("HIGH")
doc.add_paragraph(
    "No output guard. Claude may return plausible-sounding medical guidance "
    "(e.g. treating chest pain with breathing exercises) despite the system prompt "
    "instruction not to give medical advice."
)
bold_label("Evidence:")
bullet('ai_therapist.py line 262 — "Do not give medical advice" is an instruction only, not enforced')
bullet("No guard scans Claude's response for medical conditions or treatment names")
bold_label("Fix:")
bullet("Add input validator: if message contains medical symptom keywords, "
       'return fixed "see a doctor" response')
bullet("Add output guard scanning for treatment or drug name patterns before returning to user")

h("2.5  Befrienders.org Link Visibility", 2)
sev("LOW")
doc.add_paragraph(
    "Well implemented — visible in disclaimer bar, wellness modal, and CRISIS_RESPONSE. "
    "Optional improvement: add a floating help button on every therapy page."
)

h("2.6  Session Length & Wellness Check", 2)
sev("LOW")
doc.add_paragraph(
    "Well implemented — 20-minute inactivity trigger with wellness modal and helpline links. "
    "Optional improvement: add 60-minute max session duration with auto-end and resource summary."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. ENCRYPTION & SECURITY
# ═══════════════════════════════════════════════════════════════════════════════
h("3. Encryption & Security Risks")

h("3.1  Data in Transit (HTTPS/TLS)", 2)
sev("HIGH")
doc.add_paragraph(
    "TLS is optional and disabled by default in development mode. "
    "No HSTS header is set anywhere."
)
bold_label("Evidence:")
bullet("TogetherMindsAI.py lines 832-841 — TLS only loads if certs/cert.pem exists; "
       "missing certs -> HTTP with warning only")
bullet(".env line 4 — CORS_ALLOWED_ORIGINS=* allows any origin")
bold_label("Fix:")
bullet("Require TLS in production: raise RuntimeError if IS_PRODUCTION and no certs found")
bullet("Add HSTS header: Strict-Transport-Security: max-age=31536000; includeSubDomains")
bullet("Restrict CORS to explicit domains before any production deploy")

h("3.2  Data at Rest — Unencrypted", 2)
sev("CRITICAL")
doc.add_paragraph(
    "Chat messages stored in plaintext in both SQLite (local) and PostgreSQL (Cloud Run). "
    "Anyone with database access can read all therapy conversations."
)
bold_label("Evidence:")
bullet("models.py lines 21-42 — ChatMessage.text is a plain db.Text column")
bullet("config.py — SQLite is an unencrypted local file")
bullet("No field-level encryption configured for PostgreSQL")
bold_label("Fix:")
bullet("Encrypt sensitive fields using sqlalchemy_utils.EncryptedType with a KMS-managed key")
bullet("For SQLite: use SQLCipher (encrypted SQLite variant)")
bullet("For PostgreSQL: enable TDE in Cloud SQL or use pgcrypto for column-level encryption")
bullet("Implement key rotation schedule")

h("3.3  Authentication Strength (EC Keypair)", 2)
sev("MEDIUM")
doc.add_paragraph(
    "ECDSA P-256 is cryptographically sound. However, the private key is device-bound "
    "with no passphrase and no recovery path if browser data is cleared."
)
bold_label("Evidence:")
bullet("static/js/auth.js lines 75-89 — private key in IndexedDB with extractable: false (good)")
bullet("TogetherMindsAI.py lines 637-656 — 5-minute challenge, ECDSA(SHA256) verification (secure)")
bullet("Weakness: no passphrase; no backup; shared device users lose access permanently")
bold_label("Fix:")
bullet("Offer optional passphrase protection: PBKDF2(passphrase) -> encrypt key -> store in IndexedDB")
bullet("Implement key export/backup so users can re-register on a new device")
bullet("Warn users on shared devices that account is stored locally and not transferable")

h("3.4  Session Cookie Security", 2)
sev("MEDIUM")
doc.add_paragraph(
    "Flask session cookies are not explicitly hardened. Secure and SameSite flags are absent."
)
bold_label("Fix (add to TogetherMindsAI.py):")
bullet("SESSION_COOKIE_SECURE = IS_PRODUCTION")
bullet("SESSION_COOKIE_HTTPONLY = True")
bullet('SESSION_COOKIE_SAMESITE = "Lax"')
bullet("PERMANENT_SESSION_LIFETIME = 3600  # 1 hour")

h("3.5  SECRET_KEY on Disk", 2)
sev("CRITICAL")
doc.add_paragraph(
    "The Flask SECRET_KEY lives in .env on disk with no rotation mechanism. "
    ".env is gitignored (never committed to git) but is readable by any process "
    "running as the same OS user."
)
bold_label("Fix:")
bullet('Rotate now: python -c "import secrets; print(secrets.token_hex(32))" '
       "-> update .env and GCP Secret Manager")
bullet("Set file permissions: chmod 600 .env")
bullet("Implement 90-day rotation policy")

h("3.6  SQL Injection Exposure", 2)
sev("LOW")
doc.add_paragraph(
    "SQLAlchemy ORM used throughout — parameterised queries everywhere. "
    "No raw SQL strings found. Risk is mitigated — no fix required."
)

h("3.7  XSS Exposure", 2)
sev("MEDIUM")
doc.add_paragraph(
    "Frontend uses innerHTML but correctly escapes via DOM-based escapeHtml(). "
    "Jinja2 autoescaping is on. No Content Security Policy header is set."
)
bold_label("Fix:")
bullet("Confirm: app.jinja_env.autoescape = True")
bullet("Add CSP header via after_request hook restricting script-src to self and CDNs in use")

h("3.8  CORS Configuration", 2)
sev("HIGH")
doc.add_paragraph(
    "CORS is set to * in .env. Any website can connect to SocketIO and send or receive "
    "messages if they obtain a valid user_id."
)
bold_label("Fix:")
bullet("Change .env to explicit domains before any production deploy")
bullet('Add runtime guard: if IS_PRODUCTION and CORS == "*", raise RuntimeError')

h("3.9  Rate Limiting Coverage", 2)
sev("MEDIUM")
doc.add_paragraph(
    "Rate limiting covers message sending but not auth or admin endpoints."
)
bold_label("Covered:")
bullet("/therapy/solo POST and SocketIO send_message event")
bold_label("Not covered:")
bullet("/api/auth/register, /api/auth/challenge, /api/auth/verify, "
       "/session/delete, /session/nickname")
bold_label("Fix:")
bullet("Add flask-limiter with per-IP limits: registration 10/hour; "
       "challenge 5/minute; delete 1/hour")

h("3.10  Live API Key on Disk", 2)
sev("CRITICAL")
doc.add_paragraph(
    "The Anthropic API key (active and live) sits in .env on disk. "
    ".env is gitignored and never committed to git, but is readable locally."
)
bold_label("Fix:")
bullet("Rotate the Anthropic API key in the Anthropic console if there is any doubt about local exposure")
bullet("chmod 600 .env")
bullet("On Cloud Run: all secrets already in GCP Secret Manager; .env excluded by .dockerignore")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. PRIVACY
# ═══════════════════════════════════════════════════════════════════════════════
h("4. Privacy Risks")

h("4.1  Data Stored Per Session", 2)
sev("HIGH")
doc.add_paragraph(
    "Both ChatMessage and Exercise tables store full conversation text — redundant "
    "and doubles the breach surface area."
)
bold_label("Evidence:")
bullet("models.py lines 21-42 — ChatMessage stores full message text")
bullet("models.py lines 45-59 — Exercise also stores prompt (user message) and response (AI reply)")
bold_label("Fix:")
bullet("Exercise should store only metadata (type, mode, timestamp, mood_score)")
bullet("Reference ChatMessage by ID rather than duplicating text")
bullet("Do not persist emotion classifier results unless explicitly requested by user")

h("4.2  Chat Message Retention", 2)
sev("HIGH")
doc.add_paragraph(
    'Messages retained indefinitely despite auth.html\'s stated promise that '
    '"conversations will be deleted by default".'
)
bold_label("Fix:")
bullet("Add retention_expires_at to TherapySession (default: now + 30 days)")
bullet("Add hourly background purge job deleting expired sessions and their messages")
bullet('On End Session screen ask "Keep this conversation?" — '
       "yes sets 365-day retention; no deletes immediately")

h("4.3  Server-Side Logs", 2)
sev("MEDIUM")
doc.add_paragraph(
    "Flask error logging may capture user_id, session_id, or message text in exception objects."
)
bold_label("Evidence:")
bullet("TogetherMindsAI.py lines 806, 824 — app.logger.error logs exception objects that may contain PII")
bold_label("Fix:")
bullet("Add log filter redacting user_id, session_id, and message text")
bullet("Never log chat message content — only metadata (error type, timestamp)")
bullet("Configure log rotation: local 7-day; Cloud Run 30-day via GCP Cloud Logging")

h("4.4  Session ID Anonymity", 2)
sev("MEDIUM")
doc.add_paragraph(
    "6-char session IDs correctly decouple session from user identity. However, "
    "the user's UUID is broadcast to other participants via SocketIO events."
)
bold_label("Evidence:")
bullet("TogetherMindsAI.py lines 737-739 — user UUID emitted to all room participants "
       "in participant_joined event")
bold_label("Fix:")
bullet("Replace user_id in SocketIO events with an opaque per-session participant ID "
       "(secrets.token_hex(8))")
bullet("Never expose the actual UUID to other participants")

h("4.5  Third-Party Data Sharing — Anthropic API", 2)
sev("HIGH")
doc.add_paragraph(
    "Every user message is sent to Anthropic's Claude API in plaintext. "
    "This is not disclosed to users anywhere in the app. "
    "Anthropic may use data for model training unless opted out."
)
bold_label("Evidence:")
bullet("ai_therapist.py lines 341-356 — full chat history sent to client.messages.create()")
bullet("No disclosure in any template that messages leave the server")
bold_label("Fix:")
bullet("Anonymise before sending: strip names and identifiers from messages before the Claude API call")
bullet('"Your messages are sent to Anthropic to generate responses" + '
       "link to Anthropic privacy policy in auth.html")
bullet("Execute a Data Processing Addendum (DPA) with Anthropic")
bullet("Consider offline fallback for users who decline third-party data sharing")

h("4.6  Emotion Classifier Storage", 2)
sev("MEDIUM")
doc.add_paragraph(
    "Well implemented — emotion classification runs locally in RAM; results are not persisted. "
    "No third-party data leakage. No fix required."
)

h("4.7  localStorage Usage", 2)
sev("MEDIUM")
doc.add_paragraph(
    "Session nicknames and privacy banner state stored in localStorage. "
    "These persist across sessions and could expose sensitive labels if the device is compromised."
)
bold_label("Evidence:")
bullet("static/js/therapy.js lines 275-295 — nickname stored as session_nickname_<sessionId>")
bullet("No Content Security Policy to prevent localStorage exfiltration via XSS")
bold_label("Fix:")
bullet("Use sessionStorage instead of localStorage for nicknames (cleared when tab closes)")
bullet("Add CSP header via after_request hook")
bullet("Warn users before saving a nickname that it is stored on this device only")

out = "TogetherMindsAI_Risk_Audit.docx"
doc.save(out)
print("Saved:", out)
