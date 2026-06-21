import os
import config  # loads .env and exposes typed constants

# eventlet must be monkey-patched before any other imports when used
if config.ASYNC_MODE == "eventlet":
    import eventlet
    eventlet.monkey_patch()

import base64
import re
import secrets
import time
import uuid
from collections import defaultdict
from datetime import datetime, timezone, timedelta

import io
import smtplib
from email.message import EmailMessage
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, make_response, flash, send_file, abort, Response, stream_with_context
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_socketio import SocketIO, join_room, emit
from cryptography.hazmat.primitives.asymmetric.ec import ECDSA
from cryptography.hazmat.primitives.hashes import SHA256
from cryptography.hazmat.primitives.serialization import load_der_public_key
from cryptography.exceptions import InvalidSignature

from models import db, User, ChatMessage, Exercise, RateLimitEntry, TherapySession, AuditLog, Clinician, ClientAccount, SessionParticipant, NotificationLog, CopilotCard, SessionSummary, SessionHidden, SessionRecording, init_encryption
from authlib.integrations.flask_client import OAuth
from ai_therapist import detect_crisis, CRISIS_RESPONSE
import copilot
import clinical_summary
import recording
import billing
from audit import log_event
from session_id import generate_session_id, normalise_join_input, rejoin_format_hint, rejoin_placeholder
from log_filter import install_log_filter

install_log_filter()   # redact PHI-bearing fields from all log output (finding 4.3)

app = Flask(__name__)
app.config["SECRET_KEY"] = config.SECRET_KEY or os.environ.get("SECRET_KEY", "dev-fallback-key")
app.config["SQLALCHEMY_DATABASE_URI"] = config.DATABASE_URL
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = config.SQLALCHEMY_ENGINE_OPTIONS
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
# Session cookie hardening (HIPAA / finding 3.4)
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_SECURE"] = config.IS_PRODUCTION   # True on Cloud Run (HTTPS); False on localhost/test
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=30)

db.init_app(app)

# Per-IP rate limiter (finding 3.9).
# On Cloud Run, the real client IP arrives in X-Forwarded-For; using
# PROXIES_COUNT=1 tells flask-limiter to trust the first forwarded address.
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=[],                          # no global default — limits are per-route only
    storage_uri="memory://",                    # in-process; resets on restart (acceptable for rate limiting)
    headers_enabled=True,                       # send X-RateLimit-* headers so clients can back off
)
app.config["RATELIMIT_PROXIES_COUNT"] = 1       # trust one proxy hop (Cloud Run load balancer)

socketio = SocketIO(
    app,
    async_mode=config.ASYNC_MODE,
    cors_allowed_origins=config.CORS_ALLOWED_ORIGINS,
    # Disable WebSocket upgrade when running under werkzeug's dev server.
    # werkzeug does not support WebSocket; upgrade attempts cause a 500 error
    # (AssertionError: write() before start_response).  In production the
    # async_mode is eventlet, which handles WebSocket natively, so upgrades
    # are allowed there.
    allow_upgrades=not (config.FLASK_DEBUG or config.IS_TESTING),
)

_RATE_WINDOW     = config.RATE_WINDOW_SECONDS
_RATE_MAX_MSGS   = config.RATE_MAX_MESSAGES
_MAX_MSG_LEN     = config.MAX_MESSAGE_LENGTH
# Therapist-led sessions are clinical records, retained ~6 years (HIPAA § 164.312(b)).
# The expired-session purge job sweeps any row whose retention_expires_at has passed
# (e.g. legacy AI-led sessions expired by the one-time migration below).
_CLINICAL_RETENTION_DELTA = timedelta(days=2192)   # ~6 years

# ---------------------------------------------------------------------------
# Clinician OAuth (OpenID Connect) — Google & Microsoft. Registered even when
# credentials are unset (the /login buttons just won't work until they are set).
# ---------------------------------------------------------------------------
oauth = OAuth(app)
oauth.register(
    name="google",
    client_id=config.GOOGLE_CLIENT_ID,
    client_secret=config.GOOGLE_CLIENT_SECRET,
    server_metadata_url="https://accounts.google.com/.well-known/openid-configuration",
    client_kwargs={"scope": "openid"},
)
oauth.register(
    name="microsoft",
    client_id=config.MICROSOFT_CLIENT_ID,
    client_secret=config.MICROSOFT_CLIENT_SECRET,
    server_metadata_url=(
        f"https://login.microsoftonline.com/{config.MICROSOFT_TENANT}/v2.0/.well-known/openid-configuration"
    ),
    client_kwargs={"scope": "openid"},
)
_OAUTH_PROVIDERS = ("google", "microsoft")

# Azure multi-tenant ("common"/"organizations"/"consumers") publishes a discovery
# issuer templated as https://login.microsoftonline.com/{tenantid}/v2.0, but the
# real id_token substitutes the signing tenant's GUID — so Authlib's default
# "iss must equal metadata issuer" check rejects every Microsoft sign-in. Validate
# the way Microsoft documents instead: the issuer must equal the template with
# {tenantid} replaced by the token's own `tid` claim.
_MS_ISS_RE = re.compile(r"^https://login\.microsoftonline\.com/[0-9a-fA-F-]{36}/v2\.0$")


def _validate_ms_issuer(claims, value):
    """Return True if `value` is a valid Microsoft tenant issuer for this token."""
    if not value:
        return False
    tid = claims.get("tid")
    if tid:
        return value == f"https://login.microsoftonline.com/{tid}/v2.0"
    # No tid claim — fall back to accepting a well-formed tenant issuer.
    return bool(_MS_ISS_RE.match(value))



# In-memory maps (ephemeral — reset on restart, which is acceptable for these)
room_mode: dict = {}
room_participants: dict = defaultdict(set)   # session_id → set of user_ids
sid_to_user: dict = {}                        # SocketIO SID → user_id
sid_to_session: dict = {}                     # SocketIO SID → session_id

# Display name tracking — all ephemeral, reset on server restart
session_display_names: dict = {}  # session_id → {user_id: display_name}
session_taken_names: dict   = {}  # session_id → set of lowercased names (permanent for session lifetime)
session_joined_users: dict  = {}  # session_id → ordered list of user_ids (join order, for default names)

# Therapist co-pilot state — all ephemeral, repopulated on join from the DB.
session_therapist_id: dict    = {}  # session_id → therapist user_id (set only for therapist-led sessions)
session_therapist_notes: dict = {}  # session_id → list[str] of the therapist's private notes (co-pilot context)
session_recent_cards: dict    = {}  # session_id → list[str] of recently shown card texts (dedup memory)
session_friendly_name: dict   = {}  # session_id → therapist-set shared session label (ephemeral)
session_copilot_cadence: dict = {}  # session_id → "more"|"less"|"stop" live-display setting (default "more")
session_copilot_emit_counter: dict = {}  # session_id → throttle counter used by the "less" cadence

# Phase 4 recording consent state — ephemeral, reset on restart.
session_recording_requested: dict = {}                  # session_id → bool (therapist wants to record)
session_recording_consent: dict   = defaultdict(dict)   # session_id → {user_id: bool}
session_recording_active: dict     = {}                  # session_id → SessionRecording.id (or None/-1 pending)


def _record_participation(session_id: str, user_id: str) -> None:
    """Idempotently record that `user_id` joined `session_id`.

    Lets "my sessions" surface a session for a signed-in client even if they
    never spoke. Safe to call on every (re)join — the unique constraint plus the
    existence check keep it to one row per (session, user). Never raises."""
    if not session_id or not user_id:
        return
    try:
        exists = (
            db.session.query(SessionParticipant.id)
            .filter_by(session_id=session_id, user_id=user_id)
            .first()
        )
        if exists:
            return
        db.session.add(SessionParticipant(
            session_id=session_id, user_id=user_id,
            joined_at=datetime.now(timezone.utc),
        ))
        db.session.commit()
    except Exception as exc:
        # Most likely a race on the unique constraint — already recorded.
        db.session.rollback()
        app.logger.debug("participation record skipped (%s)", type(exc).__name__)


def _therapist_room(session_id: str) -> str:
    """Private SocketIO room that only the session's therapist joins.

    Co-pilot cards are emitted here, never to the main session room, so clients
    are structurally unable to receive them (they are never added to this room).
    """
    return f"{session_id}::therapist"


def _build_transcript(session_id: str, limit: int = 20) -> str:
    """Build a speaker-labelled transcript of the last `limit` messages for the co-pilot."""
    msgs = (
        ChatMessage.query
        .filter_by(session_id=session_id)
        .order_by(ChatMessage.timestamp.asc())
        .all()
    )
    therapist_id = session_therapist_id.get(session_id)
    lines = []
    for m in msgs[-limit:]:
        if m.user_id == therapist_id:
            who = "Therapist"
        elif m.user_id == "AI":
            who = "AI"
        else:
            who = m.display_name or "Client"
        lines.append(f"{who}: {m.text}")
    return "\n".join(lines)


def _run_copilot(session_id: str, mode: str, trigger_text: str = None,
                 trigger_user_id: str = None) -> None:
    """Generate co-pilot cards and emit them to the therapist-only room.

    `trigger_text`, when given, is the latest client utterance — used for the
    zero-latency keyword risk check. `trigger_user_id` is the speaker of that
    turn, stored alongside the card. Never raises into the caller.
    """
    try:
        transcript = _build_transcript(session_id)
        notes = "\n".join(session_therapist_notes.get(session_id, []))
        cards = []
        # Safety risk alerts are ALWAYS on (never gated behind a paywall).
        if trigger_text:
            cards.extend(copilot.build_risk_cards(trigger_text))
        # The AI advisory (reference + LLM suggestions) is the paid "AI analysis"
        # tier — Pro or Premium. Free clinicians still get the safety alerts above.
        if _has_ai_analysis(_session_clinician(session_id)):
            cards.extend(copilot.build_reference_cards(transcript))
            cards.extend(copilot.generate_suggestions(transcript, mode=mode, therapist_notes=notes))

        recent = session_recent_cards.setdefault(session_id, [])
        cards = copilot.dedupe_cards(cards, recent)
        if not cards:
            return
        for c in cards:
            recent.append(c["text"])
        del recent[:-30]   # cap dedup memory

        # Cards (incl. safety/risk alerts) are ALWAYS saved to the record so the
        # downloaded files are complete, regardless of the therapist's chattiness
        # setting. The setting only controls what is shown LIVE in the panel.
        _persist_cards(session_id, cards, trigger_user_id)
        log_event("suggestions_generated", session_id=session_id, count=len(cards))
        if _copilot_should_emit(session_id):
            socketio.emit("suggestion_cards", {"cards": cards}, to=_therapist_room(session_id))
    except Exception as e:
        app.logger.error("copilot error: %s", type(e).__name__)


def _copilot_should_emit(session_id: str) -> bool:
    """Whether to show co-pilot cards LIVE, per the therapist's chattiness setting:
    more = every message (default), less = ~1 in 3, stop = never. Saving to the
    record is unaffected (see _run_copilot)."""
    cadence = session_copilot_cadence.get(session_id, "more")
    if cadence == "stop":
        return False
    if cadence != "less":
        return True
    n = session_copilot_emit_counter.get(session_id, 0) + 1
    session_copilot_emit_counter[session_id] = n
    return (n % 3) == 0


def _persist_cards(session_id: str, cards: list, trigger_user_id: str = None) -> None:
    """Store emitted co-pilot cards so the console can replay full history later.

    Best-effort and isolated: a storage failure must never disrupt the live
    session (the cards have already been emitted), so it is caught and rolled
    back on its own.
    """
    import json
    try:
        for c in cards:
            db.session.add(CopilotCard(
                session_id=session_id,
                card_type=c.get("type", "observation"),
                text=c.get("text", ""),
                payload=json.dumps(c, sort_keys=True),
                confidence=c.get("confidence"),
                trigger_user_id=trigger_user_id,
            ))
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        app.logger.error("copilot card persist failed: %s", type(e).__name__)


def _emit_card_history(session_id: str) -> None:
    """Replay stored co-pilot cards to the just-joined therapist console.

    Emits to the requesting socket only (called from on_join inside the
    therapist-only branch), so clients can never receive the history.
    """
    import json
    try:
        stored = (
            CopilotCard.query
            .filter_by(session_id=session_id)
            .order_by(CopilotCard.created_at.asc(), CopilotCard.id.asc())
            .all()
        )
        if not stored:
            return
        cards = []
        for row in stored:
            try:
                cards.append(json.loads(row.payload))
            except (ValueError, TypeError):
                cards.append({"type": row.card_type, "text": row.text})
        emit("card_history", {"cards": cards})
    except Exception as e:
        app.logger.error("card history load failed: %s", type(e).__name__)


# ---------------------------------------------------------------------------
# Session modes — solo / couple / group are one multi-user realtime engine that
# differs only in presentation and the default participant-name scheme. This
# table centralises those per-mode strings; the room route and template read
# from it instead of branching. `mode` stays authoritative session data, so a
# couple session is always rendered (and advised on) as a couple.
# ---------------------------------------------------------------------------
MODE_CONFIG = {
    "solo": {
        "label": "1:1 Session",
        "icon": "person-fill",
        "placeholder": "Type a message…",
        "name_prefix": "Solo",
    },
    "couple": {
        "label": "Couple Check-in",
        "icon": "people-fill",
        "placeholder": "Type a message…",
        "name_prefix": "Partner",
    },
    "group": {
        "label": "Group Circle",
        "icon": "person-lines-fill",
        "placeholder": "Share with the group…",
        "name_prefix": "GroupMember",
    },
}


def _default_display_name(mode: str, position: int) -> str:
    """Return the default display name for a participant given their join position."""
    prefix = MODE_CONFIG.get(mode, {}).get("name_prefix", "Participant")
    return f"{prefix}{position}"


# Per-mode client capacity. "Clients" are participants other than the session's
# therapist: solo allows 1, couple allows 2, group is unlimited (no entry).
_MODE_CLIENT_CAP = {"solo": 1, "couple": 2}


def _session_is_full(session_id, mode, user_id, therapist_id) -> bool:
    """True if a non-therapist user joining would exceed the mode's client cap.

    The therapist is never counted, and the joining user is excluded from the
    tally — so the therapist, and any client reconnecting/reloading, are always
    allowed back in. Counts only participants currently present.
    """
    cap = _MODE_CLIENT_CAP.get(mode)
    if cap is None or user_id == therapist_id:
        return False
    current_clients = {
        u for u in room_participants.get(session_id, set())
        if u != therapist_id and u != user_id
    }
    return len(current_clients) >= cap


def _claim_display_name(session_id: str, user_id: str, name: str) -> None:
    """Store a display name in both active and taken dicts."""
    session_display_names.setdefault(session_id, {})[user_id] = name
    session_taken_names.setdefault(session_id, set()).add(name.lower())


def _is_name_taken(session_id: str, user_id: str, name: str) -> bool:
    """Return True if name is already taken by *another* user in this session."""
    taken = session_taken_names.get(session_id, set())
    if name.lower() not in taken:
        return False
    # Allow a user to re-confirm their own current name
    current = session_display_names.get(session_id, {}).get(user_id, "")
    return name.lower() != current.lower()


# ---------------------------------------------------------------------------
# Rate limiting — SQLite-backed sliding window
# ---------------------------------------------------------------------------

def _check_rate_limit(user_id: str) -> bool:
    """Return True if within limit, False if exceeded. Persists across restarts."""
    now = time.time()
    cutoff = now - _RATE_WINDOW
    RateLimitEntry.query.filter(
        RateLimitEntry.user_id == user_id,
        RateLimitEntry.timestamp < cutoff,
    ).delete(synchronize_session=False)
    count = RateLimitEntry.query.filter_by(user_id=user_id).count()
    if count >= _RATE_MAX_MSGS:
        db.session.commit()
        return False
    db.session.add(RateLimitEntry(user_id=user_id, timestamp=now))
    db.session.commit()
    return True


def _purge_expired_sessions():
    """Delete sessions (and their messages) whose 30-day retention window has expired.

    Also retires audit log rows older than 6 years (HIPAA § 164.312(b)).
    Called once on startup and then every 24 hours by APScheduler.
    Safe to call in tests — uses app.app_context() internally.
    """
    try:
        with app.app_context():
            now = datetime.now(timezone.utc)
            expired = TherapySession.query.filter(
                TherapySession.retention_expires_at != None,
                TherapySession.retention_expires_at < now,
            ).all()

            # Capture IDs before deletion so we can log after the commit
            purged = [(ts.id, ts.created_by) for ts in expired]

            for ts in expired:
                ChatMessage.query.filter_by(session_id=ts.id).delete(synchronize_session=False)
                CopilotCard.query.filter_by(session_id=ts.id).delete(synchronize_session=False)
                SessionSummary.query.filter_by(session_id=ts.id).delete(synchronize_session=False)
                SessionHidden.query.filter_by(session_id=ts.id).delete(synchronize_session=False)
                # DB rows only — the GCS recording objects are governed by their own
                # 30-day retention (Phase 4 Step 3).
                SessionRecording.query.filter_by(session_id=ts.id).delete(synchronize_session=False)
                Exercise.query.filter_by(user_id=ts.created_by).delete(synchronize_session=False)
                RateLimitEntry.query.filter_by(user_id=ts.created_by).delete(synchronize_session=False)
                db.session.delete(ts)

            if purged:
                db.session.commit()
                app.logger.info("Purged %d expired session(s).", len(purged))
                for sid, uid in purged:
                    log_event("session_purged_auto", session_id=sid, user_id=uid,
                              trigger="scheduler")

            # Retire audit logs older than 6 years (HIPAA minimum retention satisfied)
            six_years_ago = now - timedelta(days=6 * 365)
            old_count = AuditLog.query.filter(AuditLog.timestamp < six_years_ago).delete(
                synchronize_session=False
            )
            if old_count:
                db.session.commit()
                app.logger.info("Retired %d audit log row(s) older than 6 years.", old_count)

    except Exception as exc:
        app.logger.error("Session purge job failed: %s", exc)


if not config.IS_TESTING:
    config.secure_env_file()
    config.validate_config()
    init_encryption(config.FIELD_ENCRYPTION_KEY)
    with app.app_context():
        db.create_all()
        # One-time migration: drop the nickname column — friendly names are now
        # local-only (localStorage) and no longer stored server-side.
        from sqlalchemy import text
        try:
            db.session.execute(text("ALTER TABLE therapy_sessions DROP COLUMN nickname"))
            db.session.commit()
        except Exception:
            db.session.rollback()  # column already removed or never existed
            # rollback (not pass) is required: on Postgres a failed statement
            # aborts the transaction, which would block the migrations below.

        # One-time migration: add the therapist_id column for therapist-led
        # sessions. create_all() does not alter existing tables, so add it here.
        # Idempotent: a no-op once the column exists. Works on SQLite and Postgres.
        try:
            db.session.execute(text("ALTER TABLE therapy_sessions ADD COLUMN therapist_id VARCHAR(36)"))
            db.session.commit()
        except Exception:
            db.session.rollback()  # column already exists

    import threading   # used by the purge-job daemon thread below

    # Add retention_expires_at column if it doesn't exist yet (one-time migration)
    with app.app_context():
        from sqlalchemy import text
        try:
            db.session.execute(text(
                "ALTER TABLE therapy_sessions ADD COLUMN retention_expires_at DATETIME"
            ))
            db.session.commit()
        except Exception:
            db.session.rollback()  # column already exists (rollback clears the aborted txn)

    # Add display_name column to chat_messages for participant identification
    with app.app_context():
        from sqlalchemy import text
        try:
            db.session.execute(text(
                "ALTER TABLE chat_messages ADD COLUMN display_name VARCHAR(60)"
            ))
            db.session.commit()
        except Exception:
            db.session.rollback()  # column already exists (rollback clears the aborted txn)

    # Add the encrypted clinician email column (Phase 4 Step 3) — used to send
    # clinicians their own recording links + retention notices. Idempotent.
    with app.app_context():
        from sqlalchemy import text
        try:
            db.session.execute(text("ALTER TABLE clinicians ADD COLUMN email TEXT"))
            db.session.commit()
        except Exception:
            db.session.rollback()  # column already exists

    # Subscription billing columns on clinicians (Phase 4 Step 4). Idempotent.
    with app.app_context():
        from sqlalchemy import text
        for ddl in (
            "ALTER TABLE clinicians ADD COLUMN stripe_customer_id VARCHAR(64)",
            "ALTER TABLE clinicians ADD COLUMN plan VARCHAR(16)",
            "ALTER TABLE clinicians ADD COLUMN subscription_status VARCHAR(24)",
            # TIMESTAMP (not DATETIME): DATETIME is not a valid Postgres type, so the
            # ALTER would fail there and the column would silently never be created.
            "ALTER TABLE clinicians ADD COLUMN current_period_end TIMESTAMP",
        ):
            try:
                db.session.execute(text(ddl))
                db.session.commit()
            except Exception:
                db.session.rollback()  # column already exists

    # Add the recording 24h-before-deletion reminder timestamp (Phase 4 Step 3).
    with app.app_context():
        from sqlalchemy import text
        try:
            db.session.execute(text("ALTER TABLE session_recordings ADD COLUMN reminder_sent_at TIMESTAMP"))
            db.session.commit()
        except Exception:
            db.session.rollback()  # column already exists
        # Opaque download token (Phase 4 Step 3 security) — keeps the session id
        # out of the download URL. Idempotent.
        try:
            db.session.execute(text("ALTER TABLE session_recordings ADD COLUMN download_token VARCHAR(64)"))
            db.session.commit()
        except Exception:
            db.session.rollback()  # column already exists

    # Remove prompt/response columns from exercises — conversation text is now
    # stored only in chat_messages (encrypted). Metadata-only exercise records
    # are sufficient for the progress chart.
    with app.app_context():
        from sqlalchemy import text, inspect as sa_inspect
        col_names = {c["name"] for c in sa_inspect(db.engine).get_columns("exercises")}
        for col in ("prompt", "response"):
            if col in col_names:
                try:
                    db.session.execute(text(f"ALTER TABLE exercises DROP COLUMN {col}"))
                    db.session.commit()
                except Exception:
                    pass

    # Force-expire legacy sessions whose IDs aren't the current canonical
    # randomized-private-key format (session_id.SESSION_ID_LENGTH). The
    # scheduled purge below will sweep them in the next pass. Idempotent:
    # matching rows already have their retention bumped down, repeat runs
    # are a no-op.
    with app.app_context():
        from sqlalchemy import func
        from session_id import SESSION_ID_LENGTH
        try:
            now_utc = datetime.now(timezone.utc)
            updated = (
                TherapySession.query
                .filter(func.length(TherapySession.id) != SESSION_ID_LENGTH)
                .filter((TherapySession.retention_expires_at == None) | (TherapySession.retention_expires_at > now_utc))
                .update({TherapySession.retention_expires_at: now_utc}, synchronize_session=False)
            )
            db.session.commit()
            if updated:
                app.logger.info(
                    "Force-expired %d legacy session(s) with non-%d-char IDs.",
                    updated, SESSION_ID_LENGTH
                )
        except Exception as exc:
            db.session.rollback()
            app.logger.warning("Legacy session expiration migration skipped: %s", exc)

    from apscheduler.schedulers.background import BackgroundScheduler
    _scheduler = BackgroundScheduler()
    _scheduler.add_job(_purge_expired_sessions, "interval", hours=24, id="purge_expired_sessions")
    _scheduler.start()
    # Run once immediately on startup to catch any sessions that expired while the app was down
    threading.Thread(target=_purge_expired_sessions, daemon=True).start()
    # NOTE: the annual ICD-refresh reminder job + startup catch-up are registered
    # at the bottom of this module — they reference functions defined further down,
    # so they must be wired AFTER those defs (this block runs during import).


# ---------------------------------------------------------------------------
# Error handlers
# ---------------------------------------------------------------------------

@app.errorhandler(429)
def ratelimit_handler(e):
    """Return JSON for rate-limit errors so the client can show a clear message."""
    return jsonify({"error": "Too many attempts. Please wait a few minutes and try again."}), 429


# Routes — pages
# ---------------------------------------------------------------------------

_SW_PATH = os.path.join(os.path.dirname(__file__), "static", "js", "sw.js")


@app.route("/sw.js")
def service_worker():
    """Serve the PWA service worker with the Cloud Run revision baked in.

    Substitutes the __BUILD_VERSION__ placeholder in static/js/sw.js with the
    current K_REVISION (set automatically by Cloud Run on every deploy; falls
    back to 'dev' locally). The cache name therefore changes on every deploy,
    causing the new SW's activate handler to delete the prior deploy's cache
    and re-precache fresh assets — no manual cache-version bumps required.
    """
    with open(_SW_PATH, "r", encoding="utf-8") as f:
        body = f.read()
    build_version = os.environ.get("K_REVISION", "dev")
    body = body.replace("__BUILD_VERSION__", build_version)
    response = make_response(body)
    response.headers["Content-Type"] = "application/javascript"
    response.headers["Service-Worker-Allowed"] = "/"
    response.headers["Cache-Control"] = "no-cache"
    return response


@app.route("/.well-known/assetlinks.json")
def assetlinks():
    """Digital Asset Links — required by TWA to verify domain ownership.

    The sha256_cert_fingerprints value below is a placeholder.
    Replace it with the real SHA-256 fingerprint of your Android signing key
    after running: keytool -list -v -keystore <your.keystore>
    """
    import json as _json
    payload = [
        {
            "relation": [
                "delegate_permission/common.handle_all_urls",
                "delegate_permission/common.get_login_creds"
            ],
            "target": {
                "namespace": "android_app",
                "package_name": "com.onlinegbc.togethermindsai",
                "sha256_cert_fingerprints": [
                    "42:CD:CA:85:A9:22:D5:73:2D:CB:1D:57:25:F1:D6:48:C8:09:4C:E2:44:4D:E7:6B:EF:94:B8:BA:0A:3B",
                    "DD:6B:CA:AB:00:41:A0:1A:D4:A2:5F:F5:B5:63:28:AF:B5:06:1F:50:23:54:18:A5:87:E7:1B:D1:84:9F",
                    "DD:6B:CA:AB:00:41:A0:1A:D4:A2:5F:F5:B5:63:28:AF:B5:06:1F:50:23:54:18:A5:87:E7:1B:D1:84:9F:4B:05"
                ]
            }
        }
    ]
    response = make_response(_json.dumps(payload, indent=2))
    response.headers["Content-Type"] = "application/json"
    return response


@app.route("/")
def root():
    # Always land first-touch on the welcome/landing page.
    return redirect(url_for("welcome"))


@app.route("/welcome")
def welcome():
    return render_template("welcome.html")


# ---------------------------------------------------------------------------
# Clinician OAuth login — Google & Microsoft (OpenID Connect)
# ---------------------------------------------------------------------------

def _current_clinician_id():
    """Return the logged-in clinician's account id, or None."""
    return session.get("clinician_id")


def _current_client_account_id():
    """Return the logged-in client's account id, or None."""
    return session.get("client_account_id")


@app.context_processor
def _inject_auth_state():
    """Expose login state to every template (drives the navbar)."""
    return {
        "current_clinician_id": session.get("clinician_id"),
        "current_client_account_id": session.get("client_account_id"),
        "current_year": datetime.now(timezone.utc).year,
    }


@app.route("/login")
def login():
    """Clinician login page — Sign in with Google / Microsoft."""
    if _current_clinician_id():
        return redirect(url_for("therapist_start"))
    return render_template("login.html")


# ---------------------------------------------------------------------------
# Shared OAuth plumbing — used by BOTH the clinician routes (/auth/...) and the
# client routes (/client/auth/...). The network/OIDC exchange lives here once;
# each role's route owns only "which account to create and where to send them".
# ---------------------------------------------------------------------------

def _oauth_start(provider, callback_endpoint, scope=None):
    """Redirect the user to the provider's consent screen, returning to
    `callback_endpoint`. Returns a Flask response (redirect or error).

    `scope` overrides the registered scope for this flow only (the clinician
    flow requests "openid email"; the client flow keeps the default "openid")."""
    if provider not in _OAUTH_PROVIDERS:
        return _redirect_invalid_session()
    client = oauth.create_client(provider)
    redirect_uri = url_for(callback_endpoint, provider=provider, _external=True, _scheme="https")
    kwargs = {"scope": scope} if scope else {}
    return client.authorize_redirect(redirect_uri, **kwargs)


def _oauth_userinfo(provider):
    """Complete the OAuth exchange and return the provider's userinfo dict, or
    None on failure (the caller decides where to redirect). Never raises."""
    client = oauth.create_client(provider)
    # Microsoft multi-tenant returns a tenant-substituted issuer that fails
    # Authlib's default issuer check — validate it the way Microsoft documents.
    kwargs = {}
    if provider == "microsoft":
        kwargs["claims_options"] = {"iss": {"essential": True, "validate": _validate_ms_issuer}}
    try:
        token = client.authorize_access_token(**kwargs)
    except Exception as exc:
        app.logger.error("OAuth callback failed (%s): %s: %s", provider, type(exc).__name__, exc)
        return None
    return token.get("userinfo") or {}


def _oauth_subject(provider):
    """Complete the OAuth exchange and return the provider's stable subject id,
    or None on failure. Never raises."""
    info = _oauth_userinfo(provider)
    return info.get("sub") if info else None


def _safe_next(target):
    """Return `target` only if it's a safe same-site relative path, else None.
    Guards the post-login redirect against open-redirect abuse."""
    if target and target.startswith("/") and not target.startswith("//"):
        return target
    return None


# ---------------------------------------------------------------------------
# Clinician OAuth routes
# ---------------------------------------------------------------------------

@app.route("/auth/<provider>/login")
def oauth_login(provider):
    # Clinicians grant "email" so we can send them their own recording links +
    # retention notices (Phase 4 Step 3). The client flow stays at "openid".
    return _oauth_start(provider, "oauth_callback", scope="openid email")


@app.route("/auth/<provider>/callback")
def oauth_callback(provider):
    if provider not in _OAUTH_PROVIDERS:
        return _redirect_invalid_session()
    info = _oauth_userinfo(provider)
    subject = info.get("sub") if info else None
    if not subject:
        flash("Sign-in did not complete. Please try again.", "warning")
        return redirect(url_for("login"))
    email = (info.get("email") or "").strip().lower() or None

    now = datetime.now(timezone.utc)
    clinician = (
        Clinician.query
        .filter_by(provider=provider, provider_subject=subject)
        .first()
    )
    if clinician is None:
        clinician = Clinician(
            id=str(uuid.uuid4()), provider=provider, provider_subject=subject,
            email=email, created_at=now, last_login_at=now,
        )
        db.session.add(clinician)
        log_event("clinician_registered", user_id=clinician.id, provider=provider)
    else:
        clinician.last_login_at = now
        if email and clinician.email != email:
            clinician.email = email   # backfill / keep current for existing accounts
    db.session.commit()

    # The clinician's account id is their identity everywhere (session owner).
    session["user_id"]      = clinician.id
    session["clinician_id"] = clinician.id
    session.permanent = True
    log_event("clinician_login", user_id=clinician.id, provider=provider)
    return redirect(url_for("therapist_start"))


@app.route("/logout")
def logout():
    cid = _current_clinician_id()
    if cid:
        log_event("clinician_logout", user_id=cid)
    client_id = _current_client_account_id()
    if client_id:
        log_event("client_logout", user_id=client_id)
    session.clear()
    return redirect(url_for("welcome"))


# ---------------------------------------------------------------------------
# Subscription entitlements (Phase 4 Step 4). While BILLING_ENABLED is OFF every
# clinician keeps full access. When ON, only an active/trialing paid plan grants
# the tier: "pro" -> AI analysis, "premium" -> AI analysis + recording.
# ---------------------------------------------------------------------------

def _effective_plan(clinician):
    """The clinician's currently-entitled tier ("free"|"pro"|"premium")."""
    if not config.BILLING_ENABLED:
        return "premium"                   # billing disabled -> full access
    if clinician is None:
        return "free"
    if (clinician.subscription_status or "").lower() not in ("active", "trialing"):
        return "free"
    return clinician.plan or "free"


def _has_ai_analysis(clinician):
    return _effective_plan(clinician) in ("pro", "premium")


def _has_recording(clinician):
    return _effective_plan(clinician) == "premium"


def _session_clinician(session_id):
    """The Clinician who owns a session (its therapist), or None."""
    ts = db.session.get(TherapySession, session_id)
    if ts and ts.therapist_id:
        return db.session.get(Clinician, ts.therapist_id)
    return None


# ---------------------------------------------------------------------------
# Billing routes — Stripe Checkout + hosted billing portal + webhook. No card
# data ever reaches this app; subscription state arrives via signed webhooks.
# ---------------------------------------------------------------------------

@app.route("/billing")
def billing_page():
    cid = _current_clinician_id()
    if not cid:
        return redirect(url_for("login"))
    clin = db.session.get(Clinician, cid)
    return render_template(
        "billing.html",
        billing_enabled=config.BILLING_ENABLED,
        current_plan=(clin.plan or "free") if clin else "free",
        subscription_status=(clin.subscription_status if clin else None),
        has_customer=bool(clin and clin.stripe_customer_id),
        renews_on=(clin.current_period_end.strftime("%d %b %Y")
                   if clin and clin.current_period_end else None),
    )


@app.route("/billing/checkout/<plan>", methods=["POST"])
def billing_checkout(plan):
    cid = _current_clinician_id()
    if not cid:
        abort(403)
    if not config.BILLING_ENABLED or plan not in billing.PAID_PLANS:
        abort(404)
    clin = db.session.get(Clinician, cid)
    base = url_for("billing_page", _external=True, _scheme="https")
    url = billing.create_checkout_url(clin, plan, base + "?success=1", base + "?canceled=1")
    db.session.commit()                    # persist any newly-created stripe_customer_id
    if not url:
        flash("Could not start checkout. Please try again.", "warning")
        return redirect(url_for("billing_page"))
    log_event("billing_checkout_started", user_id=cid, plan=plan)
    return redirect(url, code=303)


@app.route("/billing/portal", methods=["POST"])
def billing_portal():
    cid = _current_clinician_id()
    if not cid:
        abort(403)
    clin = db.session.get(Clinician, cid)
    if not (clin and clin.stripe_customer_id):
        return redirect(url_for("billing_page"))
    url = billing.create_portal_url(clin.stripe_customer_id,
                                    url_for("billing_page", _external=True, _scheme="https"))
    if not url:
        flash("Could not open the billing portal. Please try again.", "warning")
        return redirect(url_for("billing_page"))
    return redirect(url, code=303)


def _clinician_for_event_object(obj):
    """Resolve the Clinician a Stripe event object belongs to — by stored customer
    id first, falling back to client_reference_id / metadata. Backfills the
    customer id when learned from checkout."""
    cust = obj.get("customer")
    clin = Clinician.query.filter_by(stripe_customer_id=cust).first() if cust else None
    if clin is None:
        ref = obj.get("client_reference_id") or (obj.get("metadata") or {}).get("clinician_id")
        if ref:
            clin = db.session.get(Clinician, ref)
            if clin and cust and not clin.stripe_customer_id:
                clin.stripe_customer_id = cust
    return clin


def _apply_checkout_completed(obj):
    clin = _clinician_for_event_object(obj)
    if clin is None:
        return
    plan = (obj.get("metadata") or {}).get("plan") or "free"
    clin.plan = plan
    clin.subscription_status = "active"
    db.session.commit()
    log_event("billing_subscribed", user_id=clin.id, plan=plan)


def _apply_subscription_change(obj):
    clin = _clinician_for_event_object(obj)
    if clin is None:
        return
    plan, status = billing.subscription_plan_and_status(obj)
    clin.plan = plan
    clin.subscription_status = status
    cpe = obj.get("current_period_end")
    if cpe:
        clin.current_period_end = datetime.fromtimestamp(cpe, tz=timezone.utc)
    db.session.commit()
    log_event("billing_subscription_updated", user_id=clin.id, plan=plan, status=status)


def _apply_subscription_deleted(obj):
    clin = _clinician_for_event_object(obj)
    if clin is None:
        return
    clin.plan = "free"
    clin.subscription_status = "canceled"
    db.session.commit()
    log_event("billing_subscription_canceled", user_id=clin.id)


@app.route("/stripe/webhook", methods=["POST"])
def stripe_webhook():
    event = billing.verify_webhook(request.get_data(), request.headers.get("Stripe-Signature", ""))
    if event is None:
        return jsonify({"error": "invalid_signature"}), 400
    etype = event.get("type", "")
    obj = (event.get("data") or {}).get("object") or {}
    try:
        if etype == "checkout.session.completed":
            _apply_checkout_completed(obj)
        elif etype in ("customer.subscription.created", "customer.subscription.updated"):
            _apply_subscription_change(obj)
        elif etype == "customer.subscription.deleted":
            _apply_subscription_deleted(obj)
    except Exception:
        db.session.rollback()
        app.logger.error("stripe webhook handling error (%s)", etype)
    return jsonify({"received": True}), 200


# ---------------------------------------------------------------------------
# Client OAuth routes — optional login so a client can find their past sessions
# across devices. Separate from the clinician routes so a client sign-in can
# never create a Clinician account; both share the _oauth_* helpers above.
# ---------------------------------------------------------------------------

@app.route("/client/login")
def client_login():
    """Optional client sign-in page — Google / Microsoft."""
    if _current_client_account_id():
        return redirect(url_for("my_sessions"))
    # Optionally remember where to return after login (e.g. a session URL).
    nxt = _safe_next(request.args.get("next"))
    if nxt:
        session["client_login_next"] = nxt
    return render_template("client_login.html")


@app.route("/client/auth/<provider>/login")
def client_oauth_login(provider):
    return _oauth_start(provider, "client_oauth_callback")


@app.route("/client/auth/<provider>/callback")
def client_oauth_callback(provider):
    if provider not in _OAUTH_PROVIDERS:
        return _redirect_invalid_session()
    subject = _oauth_subject(provider)
    if not subject:
        flash("Sign-in did not complete. Please try again.", "warning")
        return redirect(url_for("client_login"))

    now = datetime.now(timezone.utc)
    account = (
        ClientAccount.query
        .filter_by(provider=provider, provider_subject=subject)
        .first()
    )
    if account is None:
        account = ClientAccount(
            id=str(uuid.uuid4()), provider=provider, provider_subject=subject,
            created_at=now, last_login_at=now,
        )
        db.session.add(account)
        log_event("client_registered", user_id=account.id, provider=provider)
    else:
        account.last_login_at = now
    db.session.commit()

    # The account id becomes the client's stable user_id, so their messages link
    # across devices and "my sessions" can find the sessions they took part in.
    session["user_id"]           = account.id
    session["client_account_id"] = account.id
    session.permanent = True
    log_event("client_login", user_id=account.id, provider=provider)

    nxt = _safe_next(session.pop("client_login_next", None))
    return redirect(nxt or url_for("my_sessions"))


@app.route("/me/sessions")
def my_sessions():
    """A logged-in client's list of therapist-led sessions they took part in."""
    account_id = _current_client_account_id()
    if not account_id:
        return redirect(url_for("client_login"))
    # Sessions the client took part in: recorded at join time (covers silent
    # attendees) unioned with any session they sent a message in (defensive —
    # covers rows that predate join tracking).
    session_ids = {
        row[0] for row in
        db.session.query(SessionParticipant.session_id).filter_by(user_id=account_id).all()
    } | {
        row[0] for row in
        db.session.query(ChatMessage.session_id).filter_by(user_id=account_id).distinct().all()
    }
    # Drop sessions this client hid from their own view (the clinician still keeps them).
    session_ids -= {
        row[0] for row in
        db.session.query(SessionHidden.session_id).filter_by(user_id=account_id).all()
    }
    sessions = []
    if session_ids:
        sessions = (
            TherapySession.query
            .filter(TherapySession.id.in_(session_ids))
            .order_by(TherapySession.created_at.desc())
            .all()
        )
    return render_template("my_sessions.html", my_sessions=sessions)


@app.route("/therapist")
def therapist_start():
    """Landing page for a logged-in clinician to start / resume therapist-led sessions."""
    clinician_id = _current_clinician_id()
    if not clinician_id:
        return redirect(url_for("login"))
    my_sessions = (
        TherapySession.query
        .filter_by(therapist_id=clinician_id)
        .order_by(TherapySession.created_at.desc())
        .limit(20)
        .all()
    )
    return render_template("therapist.html", my_sessions=my_sessions)


@app.route("/therapist/start/<mode>", methods=["POST"])
def therapist_start_session(mode):
    """Create a new therapist-led session owned by the logged-in clinician."""
    clinician_id = _current_clinician_id()
    if not clinician_id:
        return redirect(url_for("login"))
    if mode not in ("solo", "couple", "group"):
        return _redirect_invalid_session()

    now = datetime.now(timezone.utc)
    new_sid = generate_session_id()
    db.session.add(TherapySession(
        id=new_sid, mode=mode, created_by=clinician_id,
        created_at=now,
        # Clinical record retention (~6 years), not the 30-day consumer purge.
        retention_expires_at=now + _CLINICAL_RETENTION_DELTA,
        therapist_id=clinician_id,
    ))
    db.session.commit()
    log_event("session_created", session_id=new_sid, user_id=clinician_id,
              mode=mode, therapist_led=True)
    return redirect(url_for(f"therapy_{mode}", session_id=new_sid))


@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


@app.route("/tos")
def tos():
    return render_template("tos.html")


@app.route("/auth/<therapy_mode>", methods=["GET"])
def auth_get(therapy_mode):
    # This page only exists to take a client into a clinician-led session they are
    # joining. Without a pending join there is nothing self-directed to start here.
    if not (session.get("pending_solo_session")
            or session.get("pending_couple_session")
            or session.get("pending_group_session")):
        return redirect(url_for("session_join_get"))
    return render_template("auth.html", therapy_mode=therapy_mode)


@app.route("/auth/<therapy_mode>", methods=["POST"])
@limiter.limit("10 per hour")
def auth_post(therapy_mode):
    """Legacy form-based auth fallback. Only JOINS an existing clinician-led session
    (stashed by /session/join); it never creates a new self-directed session."""
    pending_couple = session.pop("pending_couple_session", None)
    pending_group  = session.pop("pending_group_session",  None)
    pending_solo   = session.pop("pending_solo_session",   None)
    if not (pending_solo or pending_couple or pending_group):
        return redirect(url_for("session_join_get"))

    user_id = str(uuid.uuid4())
    db.session.add(User(id=user_id, therapy_mode=therapy_mode))
    db.session.commit()
    session["user_id"] = user_id

    if pending_solo:
        return redirect(url_for("therapy_solo", session_id=pending_solo))
    elif pending_couple:
        return redirect(url_for("therapy_couple", session_id=pending_couple))
    else:
        return redirect(url_for("therapy_group", session_id=pending_group))


def _session_exists(session_id: str) -> bool:
    """True iff a TherapySession row with that id exists. Used as the entry
    gate for all /therapy/<mode>/<session_id> routes — without this check,
    the routes would accept any string and silently create orphan rows."""
    if not session_id:
        return False
    return db.session.get(TherapySession, session_id) is not None


def _redirect_invalid_session():
    """Redirect to the welcome/landing page with a flash explaining the
    session is unknown/expired."""
    flash("That session doesn't exist or has expired.", "warning")
    return redirect(url_for("welcome"))


def _render_session_room(session_id, mode):
    """Render the shared realtime room for any mode (solo / couple / group).

    The three legacy routes are thin wrappers over this. `ts.mode` is the
    authoritative mode — a URL pointing at the wrong one is redirected to the
    canonical room, so a couple session is always shown as a couple.
    """
    user_id = session.get("user_id")
    if not user_id:
        return redirect(url_for("auth_get", therapy_mode=mode))
    if not _session_exists(session_id):
        return _redirect_invalid_session()
    ts = db.session.get(TherapySession, session_id)

    # The stored mode wins over the URL — keeps every mode honestly labelled.
    if ts and ts.mode in MODE_CONFIG and ts.mode != mode:
        return redirect(url_for(f"therapy_{ts.mode}", session_id=session_id))

    # Solo is therapist-led only on this branch; a solo room without a therapist
    # is a stale/invalid record (the consumer 1:1 flow was removed).
    if mode == "solo" and not (ts and ts.therapist_id):
        return _redirect_invalid_session()

    # Enforce the per-mode client capacity before admitting a client into the room.
    if _session_is_full(session_id, mode, user_id, ts.therapist_id if ts else None):
        flash("That session is already full.", "warning")
        return redirect(url_for("welcome"))

    cfg = MODE_CONFIG.get(mode, MODE_CONFIG["solo"])
    return render_template(
        "session_live.html",
        user_id=user_id, session_id=session_id, mode=mode,
        mode_label=cfg["label"], mode_icon=cfg["icon"], mode_placeholder=cfg["placeholder"],
        is_therapist=bool(ts and ts.therapist_id and ts.therapist_id == user_id),
        is_therapist_led=bool(ts and ts.therapist_id),
        rtc_enabled=config.RTC_ENABLED,
        recording_enabled=config.RECORDING_ENABLED,
    )


@app.route("/therapy/solo/<session_id>", methods=["GET"])
def therapy_solo(session_id):
    return _render_session_room(session_id, "solo")


@app.route("/therapy/couple/<session_id>")
def therapy_couple(session_id):
    return _render_session_room(session_id, "couple")


@app.route("/therapy/group/<session_id>")
def therapy_group(session_id):
    return _render_session_room(session_id, "group")


# ---------------------------------------------------------------------------
# Realtime conferencing (Phase 1) — short-lived tokens minted server-side so the
# LiveKit secret and AssemblyAI key never reach the browser. A caller must have a
# session identity and be pointing at a real session.
# ---------------------------------------------------------------------------

def _rtc_guard(session_id):
    """Return (user_id, error_response). error_response is None when allowed."""
    if not config.RTC_ENABLED:
        return None, (jsonify({"error": "rtc_disabled"}), 503)
    user_id = session.get("user_id")
    if not user_id:
        return None, (jsonify({"error": "no_identity"}), 403)
    if not _session_exists(session_id):
        return None, (jsonify({"error": "no_session"}), 404)
    return user_id, None


@app.route("/rtc/livekit-token", methods=["POST"])
@limiter.limit("30 per minute")
def rtc_livekit_token():
    """Mint a LiveKit join token for the caller to join the session's audio room."""
    session_id = (request.get_json(silent=True) or {}).get("session_id", "")
    user_id, err = _rtc_guard(session_id)
    if err:
        return err
    from livekit.api import AccessToken, VideoGrants
    display = session_display_names.get(session_id, {}).get(user_id) or "Participant"
    token = (
        AccessToken(config.LIVEKIT_API_KEY, config.LIVEKIT_API_SECRET)
        .with_identity(user_id)
        .with_name(display)
        .with_grants(VideoGrants(
            room_join=True, room=session_id,
            can_publish=True, can_subscribe=True,
        ))
        .to_jwt()
    )
    return jsonify({"token": token, "url": config.LIVEKIT_URL})


@app.route("/rtc/stt-token", methods=["POST"])
@limiter.limit("30 per minute")
def rtc_stt_token():
    """Mint a short-lived AssemblyAI streaming token for in-browser transcription."""
    session_id = (request.get_json(silent=True) or {}).get("session_id", "")
    user_id, err = _rtc_guard(session_id)
    if err:
        return err
    import requests
    try:
        resp = requests.get(
            "https://streaming.assemblyai.com/v3/token",
            headers={"Authorization": config.ASSEMBLYAI_API_KEY},
            params={"expires_in_seconds": 600},
            timeout=10,
        )
        resp.raise_for_status()
        return jsonify({"token": resp.json().get("token")})
    except Exception as exc:
        app.logger.error("AssemblyAI token error: %s", type(exc).__name__)
        return jsonify({"error": "stt_token_failed"}), 502


@app.route("/api/display-name", methods=["POST"])
def api_set_display_name():
    """AJAX endpoint used by solo mode to set or rename a display name.

    Couple/group modes use the set_display_name / rename socket events instead.
    Request JSON: {session_id, display_name}
    Response JSON: {display_name} on success, {error} on conflict/validation error.
    """
    user_id = session.get("user_id")
    if not user_id:
        return jsonify({"error": "Not authenticated"}), 401

    data = request.get_json(silent=True) or {}
    session_id = data.get("session_id", "").strip()
    name = data.get("display_name", "").strip()

    if not name:
        return jsonify({"error": "Display name cannot be empty."}), 400
    if len(name) > 40:
        return jsonify({"error": "Display name must be 40 characters or fewer."}), 400

    if _is_name_taken(session_id, user_id, name):
        return jsonify({"error": f"'{name}' is already taken in this session. Please choose another."}), 409

    _claim_display_name(session_id, user_id, name)
    return jsonify({"display_name": name}), 200


@app.route("/api/translate-check", methods=["POST"])
@limiter.limit("120 per hour")
def api_translate_check():
    """Detect language; if non-English, translate to English.

    Returns:
      {"is_english": true} — text is English; client sends as-is
      {"is_english": false, "translation": "<English>"} — show user the
        translation and have them confirm before the message proceeds

    On any error: returns is_english=true so the user is never blocked.
    Crisis-signal preservation matters: uses Sonnet 4.6 with a prompt that
    explicitly preserves emotional intensity and clinical idioms (e.g.,
    "I want to disappear" must keep its suicidal-ideation connotation in the
    English output, not flatten to "I want to leave").
    """
    user_id = session.get("user_id")
    if not user_id:
        return jsonify({"error": "Not authenticated"}), 401

    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "Empty text"}), 400
    if len(text) > config.MAX_MESSAGE_LENGTH:
        return jsonify({"error": "Text too long"}), 413

    try:
        from ai_therapist import _get_claude_client
        client = _get_claude_client()
        result = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=600,
            messages=[{
                "role": "user",
                "content": (
                    "Detect the language of the message below.\n"
                    "If it is in English, reply with EXACTLY the two letters: EN\n"
                    "Otherwise, reply with ONLY the English translation. "
                    "Preserve emotional intensity and any culturally specific idioms "
                    "in their clinical sense (for example, '消えたい' or "
                    "'me quiero morir' must convey suicidal ideation, not be flattened "
                    "to 'I want to leave'). Do not add explanations, prefixes, or quotes.\n\n"
                    "Message:\n" + text
                )
            }]
        )
        reply = (result.content[0].text or "").strip()
        if reply.upper() == "EN":
            return jsonify({"is_english": True}), 200
        return jsonify({"is_english": False, "translation": reply}), 200
    except Exception as exc:
        app.logger.error("translate-check failed for user %s: %s", user_id, exc)
        # Fail open — never block the user from sending due to a translate error.
        return jsonify({"is_english": True, "fallback": True}), 200


# ---------------------------------------------------------------------------
# Feedback form — email-only delivery, no DB storage, no audit log entry,
# no IP / user_id / PII captured. Per FEEDBACK_FORM_PLAN.md.
# ---------------------------------------------------------------------------

_FEEDBACK_TEXT_MAX = 1000
_FEEDBACK_PLATFORMS = {"web", "android_twa", "ios_pwa", "mobile_browser"}
_FEEDBACK_MODES = {"solo", "couple", "group", None}
_FEEDBACK_PAY = {"yes", "maybe", "no", None}
_FEEDBACK_RATINGS = {1, 2, 3, 4, 5, None}
_FEEDBACK_OS = {"windows", "macos", "linux", "android", "ios", "unknown", None}

# In-memory cooldown keyed by Flask session cookie (NOT IP).
# Resets on process restart — acceptable for an anti-spam guard.
_feedback_last_submit: dict = {}   # session_key -> last unix epoch
_feedback_daily_count: dict = {}   # (session_key, utc_date_iso) -> count

_FEEDBACK_COOLDOWN_SECONDS = 60
_FEEDBACK_DAILY_MAX = 10


def _feedback_session_key() -> str:
    """Stable per-session key used for rate limiting. Creates a server-side
    token in the Flask session if one is missing — no IP, no user_id."""
    key = session.get("_fb_key")
    if not key:
        key = secrets.token_urlsafe(16)
        session["_fb_key"] = key
    return key


def _device_label(platform: str, os_name) -> str:
    """Friendly human-readable device descriptor combining shell + OS."""
    if platform == "android_twa":
        return "Android (installed app)"
    if platform == "ios_pwa":
        return "iPhone/iPad (installed app)"
    if platform == "mobile_browser":
        if os_name == "android":
            return "Android (mobile browser)"
        if os_name == "ios":
            return "iPhone/iPad (mobile browser)"
        return "Mobile browser"
    # platform == "web"
    if os_name == "windows":
        return "Windows laptop / desktop"
    if os_name == "macos":
        return "Mac laptop / desktop"
    if os_name == "linux":
        return "Linux laptop / desktop"
    return "Laptop / desktop"


def _mode_label(mode) -> str:
    return {
        "solo": "1:1 Session",
        "couple": "Couple Check-in",
        "group": "Group Circle",
    }.get(mode, "Not in a session")


def _pay_label(pay) -> str:
    return {"yes": "Yes", "maybe": "Maybe", "no": "No"}.get(pay, "Not answered")


def _stars(rating) -> str:
    if not rating:
        return "Not rated"
    filled = "★" * int(rating)
    empty = "☆" * (5 - int(rating))
    return f"{filled}{empty}"


def _format_feedback_email(payload: dict) -> tuple:
    """Build (subject, plain_body, html_body) for the feedback email."""
    rating = payload.get("rating")
    rating_str = f"{rating} / 5" if rating else "N/A"
    pay = payload.get("would_pay")
    platform = payload.get("platform") or ""
    os_name = payload.get("os")
    mode = payload.get("mode")
    device = _device_label(platform, os_name)
    mode_label = _mode_label(mode)
    pay_label = _pay_label(pay)
    stars = _stars(rating)
    submitted_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    subject = f"TogetherMindsAI feedback — {mode_label} on {device} — {rating_str}"

    # ----- Plain-text body (fallback) ---------------------------------------
    def text_section(title: str, text: str) -> str:
        text = (text or "").strip()
        if not text:
            return f"{title}:\n  (none)\n"
        return f"{title}:\n{text}\n"

    plain = (
        f"Rating:           {rating_str}\n"
        f"Would pay:        {pay_label}\n"
        f"Device:           {device}\n"
        f"Session mode:     {mode_label}\n"
        f"Submitted at:     {submitted_at}\n"
        f"\n"
        + text_section("What worked well", payload.get("what_worked", ""))
        + "\n"
        + text_section("What could be improved", payload.get("what_to_improve", ""))
        + "\n"
        + text_section("Desired features", payload.get("desired_features", ""))
        + "\n"
        + text_section("Anything else", payload.get("other", ""))
    )

    # ----- HTML body --------------------------------------------------------
    from html import escape as h

    def html_card(title: str, text: str) -> str:
        text = (text or "").strip()
        if not text:
            return (
                f'<h3 style="color:#388E3C;font-size:15px;margin:24px 0 8px;font-weight:600;">{h(title)}</h3>'
                f'<div style="padding:14px 16px;background:#f7faf7;border-left:3px solid #d0d0d0;border-radius:4px;color:#999;font-style:italic;">(none)</div>'
            )
        return (
            f'<h3 style="color:#388E3C;font-size:15px;margin:24px 0 8px;font-weight:600;">{h(title)}</h3>'
            f'<div style="padding:14px 16px;background:#f7faf7;border-left:3px solid #4CAF50;border-radius:4px;white-space:pre-wrap;line-height:1.5;">{h(text)}</div>'
        )

    rating_html = (
        f'<span style="color:#FFC107;font-size:18px;letter-spacing:2px;">{stars}</span> '
        f'<span style="color:#999;margin-left:6px;">({h(rating_str)})</span>'
    ) if rating else f'<span style="color:#999;font-style:italic;">{h(stars)}</span>'

    meta_row = (
        '<tr>'
        '<td style="padding:14px 16px;border-bottom:1px solid #e6efe6;width:160px;font-weight:600;color:#555;">{label}</td>'
        '<td style="padding:14px 16px;border-bottom:1px solid #e6efe6;color:#212121;">{value}</td>'
        '</tr>'
    )
    meta_row_last = (
        '<tr>'
        '<td style="padding:14px 16px;width:160px;font-weight:600;color:#555;">{label}</td>'
        '<td style="padding:14px 16px;color:#212121;">{value}</td>'
        '</tr>'
    )

    html_body = f"""<!DOCTYPE html>
<html>
<body style="margin:0;padding:0;background:#f5f7f5;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;color:#212121;">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f5f7f5;padding:24px 0;">
<tr><td align="center">
<table width="600" cellpadding="0" cellspacing="0" style="max-width:600px;width:100%;background:#ffffff;border-radius:12px;overflow:hidden;box-shadow:0 2px 6px rgba(0,0,0,0.06);">
  <tr><td style="background:#4CAF50;padding:22px 28px;color:#ffffff;">
    <div style="font-size:12px;letter-spacing:0.08em;text-transform:uppercase;opacity:0.9;">TogetherMindsAI</div>
    <div style="font-size:22px;font-weight:600;margin-top:4px;">New feedback received</div>
  </td></tr>

  <tr><td style="padding:24px 28px 8px;">
    <table width="100%" cellpadding="0" cellspacing="0" style="background:#f7faf7;border-radius:8px;border:1px solid #e6efe6;border-collapse:separate;border-spacing:0;">
      {meta_row.format(label='Rating', value=rating_html)}
      {meta_row.format(label='Would pay', value=h(pay_label))}
      {meta_row.format(label='Device', value=h(device))}
      {meta_row.format(label='Session mode', value=h(mode_label))}
      {meta_row_last.format(label='Submitted', value=h(submitted_at))}
    </table>
  </td></tr>

  <tr><td style="padding:0 28px 24px;">
    {html_card('What worked well', payload.get('what_worked', ''))}
    {html_card('What could be improved', payload.get('what_to_improve', ''))}
    {html_card('Desired features', payload.get('desired_features', ''))}
    {html_card('Anything else', payload.get('other', ''))}
  </td></tr>

  <tr><td style="padding:14px 24px;background:#fafafa;color:#999;font-size:11px;text-align:center;border-top:1px solid #eee;">
    No IP, no name, no session content captured.
  </td></tr>
</table>
</td></tr>
</table>
</body>
</html>"""

    return subject, plain, html_body


def _send_feedback_email(subject: str, plain_body: str, html_body: str) -> None:
    """Send feedback via Gmail SMTP as multipart/alternative (text + HTML).
    Raises on failure — caller maps to 503."""
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = config.FEEDBACK_FROM_EMAIL or config.FEEDBACK_SMTP_USER
    msg["To"] = ", ".join(config.FEEDBACK_TO_EMAILS)
    msg.set_content(plain_body)
    msg.add_alternative(html_body, subtype="html")

    with smtplib.SMTP(config.FEEDBACK_SMTP_HOST, config.FEEDBACK_SMTP_PORT, timeout=5) as smtp:
        smtp.starttls()
        smtp.login(config.FEEDBACK_SMTP_USER, config.FEEDBACK_SMTP_PASSWORD)
        smtp.send_message(msg)


def _send_email(to_emails, subject: str, plain_body: str, html_body: str) -> None:
    """Send a multipart/alternative email to an arbitrary recipient list via the
    same Gmail SMTP path. Raises on failure (callers decide how to handle)."""
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = config.FEEDBACK_FROM_EMAIL or config.FEEDBACK_SMTP_USER
    msg["To"] = ", ".join(to_emails)
    msg.set_content(plain_body)
    msg.add_alternative(html_body, subtype="html")

    with smtplib.SMTP(config.FEEDBACK_SMTP_HOST, config.FEEDBACK_SMTP_PORT, timeout=5) as smtp:
        smtp.starttls()
        smtp.login(config.FEEDBACK_SMTP_USER, config.FEEDBACK_SMTP_PASSWORD)
        smtp.send_message(msg)


# ---------------------------------------------------------------------------
# Annual ICD code-refresh reminder — emailed from the app every March 1.
# WHO ships ~annual ICD-11 releases; this nudges a refresh of the therapist
# co-pilot's ICD grounding (clinical_data/icd_corpus.json) and carries the full
# runbook so the steps don't have to be remembered. Reuses the feedback SMTP
# path and admin inbox (config.FEEDBACK_TO_EMAILS).
# ---------------------------------------------------------------------------

def _icd_refresh_reminder_content():
    """Return (subject, plain_body, html_body) for the annual ICD-refresh email."""
    subject = "TogetherMindsAI — Annual ICD code refresh due (co-pilot grounding)"

    steps = [
        ("Why", "WHO publishes ICD-11 on a roughly annual release cycle. Refresh the therapist "
                "co-pilot's ICD grounding so the ICD-11 deep links point at the current release."),
        ("1. Credentials", "The WHO ICD-API is free (registration only). Client credentials are "
                "WHO_ICD_ClientId / WHO_ICD_ClientSecret, stored in the local .env (gitignored) AND "
                "in Google Secret Manager. Make sure .env has them before running the script."),
        ("2. Harvest (dry run)", "Run: python scripts/harvest_icd11_entities.py — this authenticates "
                "to the ICD-API, prints the latest MMS release, and resolves each ICD-11 code in "
                "clinical_data/icd_corpus.json to its WHO Foundation entity id + deep-link URL."),
        ("3. Apply", "Re-run with --write: python scripts/harvest_icd11_entities.py --write — this "
                "writes the refreshed icd11_url deep links back into clinical_data/icd_corpus.json."),
        ("4. Verify", "Spot-check a few resolved entity titles match the corpus labels, then run: "
                "pytest tests/test_clinical_reference.py tests/test_copilot.py"),
        ("5. Commit", "Commit & push the updated clinical_data/icd_corpus.json (never commit .env)."),
        ("ICD-10 note", "ICD-10 links are search-by-code and need no refresh — they don't break on a "
                "new release. The ICD-11 deep links are the thing this reminder is about."),
    ]

    plain = (
        "Annual ICD code refresh for the therapist co-pilot.\n\n"
        + "\n\n".join(f"{label}:\n  {text}" for label, text in steps)
        + "\n\nSee project memory: project_who_icd_api.\n"
    )

    from html import escape as h
    rows = "".join(
        f'<h3 style="color:#388E3C;font-size:15px;margin:18px 0 4px;">{h(label)}</h3>'
        f'<div style="line-height:1.5;color:#212121;">{h(text)}</div>'
        for label, text in steps
    )
    html_body = (
        '<div style="font-family:-apple-system,Segoe UI,Roboto,sans-serif;max-width:640px;">'
        '<h2 style="color:#2e7d32;">Annual ICD code refresh due</h2>'
        '<p style="color:#555;">A yearly nudge to refresh the therapist co-pilot\'s ICD grounding.</p>'
        f"{rows}"
        '<p style="color:#888;font-size:12px;margin-top:20px;">Sent automatically by TogetherMindsAI on March 1.</p>'
        "</div>"
    )
    return subject, plain, html_body


_ICD_REMINDER_KEY = "icd_refresh"
_ICD_REMINDER_START_YEAR = 2027


def _claim_notification(key: str, year: int) -> bool:
    """Atomically claim (key, year) in the notification ledger.

    Returns True iff this caller inserted the row — i.e. won the right to send.
    The unique (key, year) constraint makes this exactly-once across instances
    and restarts. Assumes an active app context.
    """
    from sqlalchemy.exc import IntegrityError
    try:
        db.session.add(NotificationLog(key=key, year=year))
        db.session.commit()
        return True
    except IntegrityError:
        db.session.rollback()          # someone else already claimed this year
        return False


def _release_notification(key: str, year: int) -> None:
    """Release a previously-claimed (key, year) so a later run can retry."""
    try:
        NotificationLog.query.filter_by(key=key, year=year).delete()
        db.session.commit()
    except Exception:
        db.session.rollback()


def _deliver_icd_reminder(year: int) -> None:
    """Send the ICD-refresh reminder for `year` exactly once. Never raises.

    Claims the year first; on send failure the claim is released so a later
    startup/cron run retries rather than silently losing the year.
    """
    if year < _ICD_REMINDER_START_YEAR:
        return
    if not (config.FEEDBACK_SMTP_USER and config.FEEDBACK_SMTP_PASSWORD):
        app.logger.warning("ICD refresh reminder skipped — SMTP creds not configured.")
        return
    with app.app_context():
        if not _claim_notification(_ICD_REMINDER_KEY, year):
            return                     # already sent for this year
        try:
            subject, plain, html_body = _icd_refresh_reminder_content()
            _send_feedback_email(subject, plain, html_body)
            app.logger.info("ICD refresh reminder email sent for %d.", year)
        except Exception:
            # Don't log the body — it could echo SMTP server detail.
            _release_notification(_ICD_REMINDER_KEY, year)
            app.logger.warning("ICD refresh reminder email failed; claim released for retry.")


def _send_icd_refresh_reminder() -> None:
    """Cron target — fires March 1 (UTC). Delegates to the exactly-once delivery."""
    _deliver_icd_reminder(datetime.now(timezone.utc).year)


def _icd_reminder_catchup() -> None:
    """Startup catch-up: if we're on/after March 1 and this year's reminder has
    not gone out yet, send it now (late is better than never). The (key, year)
    claim guarantees it won't duplicate the cron send or another instance."""
    now = datetime.now(timezone.utc)
    if (now.month, now.day) >= (3, 1):
        _deliver_icd_reminder(now.year)


# Wire the annual ICD-refresh reminder now that its functions are defined. This
# runs during import AFTER the main startup block above, so `_scheduler` and
# `threading` (both module globals created there) already exist. Keeping these
# two references below the defs avoids the import-order NameError that would
# occur if they were registered up in the main startup block.
if not config.IS_TESTING:
    # Annual reminder, emailed to the admin inbox every March 1 (UTC).
    _scheduler.add_job(
        _send_icd_refresh_reminder, "cron",
        month=3, day=1, hour=9, timezone="UTC", id="icd_refresh_reminder",
    )
    # Catch-up if the app was scaled to zero / down on March 1 (exactly-once via ledger).
    threading.Thread(target=_icd_reminder_catchup, daemon=True).start()


@app.route("/feedback")
def feedback_page():
    """Standalone feedback form. Public — no login required."""
    return render_template("feedback.html")


@app.route("/api/feedback", methods=["POST"])
def api_feedback():
    """Receive feedback, send via email. Stateless — no DB, no audit, no IP."""
    payload = request.get_json(silent=True) or {}

    # --- Validation ---------------------------------------------------------
    rating = payload.get("rating")
    if rating is not None:
        # bool is a subclass of int — reject explicitly. Reject floats and
        # strings outright so the contract is "rating is int|null".
        if isinstance(rating, bool) or not isinstance(rating, int):
            return jsonify({"error": "Invalid rating"}), 400
    if rating not in _FEEDBACK_RATINGS:
        return jsonify({"error": "Invalid rating"}), 400

    would_pay = payload.get("would_pay")
    if would_pay == "":
        would_pay = None
    if would_pay is not None and not isinstance(would_pay, str):
        return jsonify({"error": "Invalid would_pay"}), 400
    if would_pay not in _FEEDBACK_PAY:
        return jsonify({"error": "Invalid would_pay"}), 400

    platform = payload.get("platform")
    if not isinstance(platform, str) or platform not in _FEEDBACK_PLATFORMS:
        return jsonify({"error": "Invalid platform"}), 400

    os_name = payload.get("os")
    if os_name == "":
        os_name = None
    if os_name is not None and not isinstance(os_name, str):
        return jsonify({"error": "Invalid os"}), 400
    if os_name not in _FEEDBACK_OS:
        return jsonify({"error": "Invalid os"}), 400

    mode = payload.get("mode")
    if mode == "":
        mode = None
    if mode is not None and not isinstance(mode, str):
        return jsonify({"error": "Invalid mode"}), 400
    if mode not in _FEEDBACK_MODES:
        return jsonify({"error": "Invalid mode"}), 400

    text_fields = ("what_worked", "what_to_improve", "desired_features", "other")
    cleaned = {}
    for f in text_fields:
        val = payload.get(f, "") or ""
        if not isinstance(val, str):
            return jsonify({"error": f"Invalid {f}"}), 400
        if len(val) > _FEEDBACK_TEXT_MAX:
            return jsonify({"error": f"{f} too long (max {_FEEDBACK_TEXT_MAX} chars)"}), 400
        cleaned[f] = val.strip()

    has_content = (
        rating is not None
        or would_pay is not None
        or any(cleaned[f] for f in text_fields)
    )
    if not has_content:
        return jsonify({"error": "Empty submission"}), 400

    # --- Rate limiting (cookie-based, no IP) --------------------------------
    key = _feedback_session_key()
    now = time.time()
    today_iso = datetime.now(timezone.utc).date().isoformat()

    last = _feedback_last_submit.get(key, 0)
    if now - last < _FEEDBACK_COOLDOWN_SECONDS:
        return jsonify({"error": "Please wait a moment before submitting again."}), 429

    daily_key = (key, today_iso)
    # Drop stale daily entries (different date)
    stale_keys = [k for k in _feedback_daily_count if k[0] == key and k[1] != today_iso]
    for k in stale_keys:
        _feedback_daily_count.pop(k, None)
    if _feedback_daily_count.get(daily_key, 0) >= _FEEDBACK_DAILY_MAX:
        return jsonify({"error": "Daily feedback limit reached."}), 429

    # --- Build & send -------------------------------------------------------
    subject, plain_body, html_body = _format_feedback_email({
        "rating": rating,
        "would_pay": would_pay,
        "platform": platform,
        "os": os_name,
        "mode": mode,
        **cleaned,
    })

    if not (config.FEEDBACK_SMTP_USER and config.FEEDBACK_SMTP_PASSWORD):
        return jsonify({"error": "Feedback service not configured"}), 503

    try:
        _send_feedback_email(subject, plain_body, html_body)
    except Exception:
        # Intentionally do not log the exception body — it could echo SMTP
        # auth or the email content. A generic warning suffices.
        app.logger.warning("Feedback SMTP send failed")
        return jsonify({"error": "Could not send feedback right now. Please try again later."}), 503

    _feedback_last_submit[key] = now
    _feedback_daily_count[daily_key] = _feedback_daily_count.get(daily_key, 0) + 1
    return jsonify({"ok": True}), 200


@app.route("/progress/<user_id>/<therapy_mode>")
def progress(user_id, therapy_mode):
    # Access control: a user may only view their OWN progress page. Without this,
    # the page (which exposes the session ID + activity data + download links) was
    # an IDOR — any user_id in the URL would render it.
    if session.get("user_id") != user_id:
        abort(403)
    exercises = (
        Exercise.query
        .filter_by(user_id=user_id)
        .order_by(Exercise.timestamp.asc())
        .all()
    )

    week_counts = {}
    for ex in exercises:
        iso = ex.timestamp.isocalendar()
        week_label = f"{iso[0]}-W{iso[1]:02d}"
        week_counts[week_label] = week_counts.get(week_label, 0) + 1

    chart_data = [{"week": w, "count": c} for w, c in sorted(week_counts.items())]

    # Find the user's most recent session (the session_id of their latest
    # message). Used for the transcript download links — without this the
    # template would have to guess which session to download.
    latest_session_id = (
        db.session.query(ChatMessage.session_id)
        .filter_by(user_id=user_id)
        .order_by(ChatMessage.timestamp.desc())
        .limit(1)
        .scalar()
    )

    return render_template(
        "progress.html",
        chart_data=chart_data,
        user_id=user_id,
        therapy_mode=therapy_mode,
        session_id=latest_session_id,
        # The clinician owns the clinical record, so "Hide my data" is a client-only
        # control — never shown to a logged-in clinician viewing their own progress.
        is_clinician=bool(session.get("clinician_id")),
    )


# ---------------------------------------------------------------------------
# Routes — session resumption
# ---------------------------------------------------------------------------

def _join_template(**kwargs):
    """Render join_session.html with format hint context always populated."""
    return render_template(
        "join_session.html",
        rejoin_hint=rejoin_format_hint(),
        rejoin_placeholder=rejoin_placeholder(),
        **kwargs,
    )


@app.route("/session/join", methods=["GET"])
def session_join_get():
    return _join_template()


@app.route("/session/join", methods=["POST"])
def session_join_post():
    raw = request.form.get("session_id", "").strip()
    if not raw:
        return _join_template(error="Please enter a Session ID.")

    # Case-insensitive lookup: normalize both sides to uppercase
    from sqlalchemy import func as sa_func
    ts = TherapySession.query.filter(
        sa_func.upper(TherapySession.id) == raw.upper()
    ).first()
    if not ts:
        return _join_template(error="Session not found. Check the ID and try again.")

    session_id = ts.id  # always use the real ID from DB

    if ts.mode == "solo":
        # Solo sessions are therapist-led only on this branch; a solo row without a
        # therapist_id is a stale/invalid record and is treated as not found.
        if not ts.therapist_id:
            return _join_template(error="Session not found. Check the ID and try again.")
        # Therapist-led 1:1 — the joiner is a CLIENT and must get their OWN identity
        # (never the therapist's). Mirror the couple/group join flow.
        user_id = session.get("user_id")
        if not user_id:
            session["pending_solo_session"] = session_id
            return redirect(url_for("auth_get", therapy_mode="solo"))
        return redirect(url_for("therapy_solo", session_id=session_id))

    user_id = session.get("user_id")
    if not user_id:
        # Stash the real session_id so auth_post can redirect back to the right room
        if ts.mode == "couple":
            session["pending_couple_session"] = session_id
        else:
            session["pending_group_session"] = session_id
        return redirect(url_for("auth_get", therapy_mode=ts.mode))

    if ts.mode == "couple":
        return redirect(url_for("therapy_couple", session_id=session_id))
    else:
        return redirect(url_for("therapy_group", session_id=session_id))




# ---------------------------------------------------------------------------
# Routes — GDPR data deletion
# ---------------------------------------------------------------------------

@app.route("/user/<user_id>", methods=["DELETE"])
def delete_user(user_id):
    """Handle a user's "delete my data" request.

    A therapist-led session is a CLINICAL RECORD the clinician must retain under
    medical-record law — neither participant may erase it. So for those sessions
    we HIDE them from the requesting user's own view (the clinician's copy is
    untouched). Genuinely non-clinical data (consumer sessions, exercises, the
    account row when there is no clinical footprint) is still permanently erased.
    """
    if session.get("user_id") != user_id:
        return jsonify({"error": "Forbidden"}), 403

    affected = {
        sid for (sid,) in db.session.query(ChatMessage.session_id)
        .filter_by(user_id=user_id).distinct()
    } | {
        sid for (sid,) in db.session.query(SessionParticipant.session_id)
        .filter_by(user_id=user_id).distinct()
    }

    hidden_clinical, erasable, owns_clinical = [], [], False
    for sid in affected:
        ts = db.session.get(TherapySession, sid)
        if ts and ts.therapist_id:                 # clinical record — must be retained
            if ts.therapist_id == user_id:
                owns_clinical = True                # the clinician: retain, untouched
            else:
                hidden_clinical.append(sid)         # a client: hide from their own view
        else:
            erasable.append(sid)                    # non-clinical — may be erased

    now = datetime.now(timezone.utc)
    for sid in hidden_clinical:
        if not db.session.query(SessionHidden.id).filter_by(session_id=sid, user_id=user_id).first():
            db.session.add(SessionHidden(session_id=sid, user_id=user_id, hidden_at=now))

    for sid in erasable:
        ChatMessage.query.filter_by(session_id=sid, user_id=user_id).delete(synchronize_session=False)
        SessionParticipant.query.filter_by(session_id=sid, user_id=user_id).delete(synchronize_session=False)

    # Non-clinical account metadata is always erasable.
    Exercise.query.filter_by(user_id=user_id).delete(synchronize_session=False)
    RateLimitEntry.query.filter_by(user_id=user_id).delete(synchronize_session=False)

    # Fully remove the account only when the user has NO clinical footprint at all
    # (their identity must persist while they appear in a retained clinical record).
    has_clinical = bool(hidden_clinical) or owns_clinical
    if not has_clinical:
        SessionParticipant.query.filter_by(user_id=user_id).delete(synchronize_session=False)
        User.query.filter_by(id=user_id).delete(synchronize_session=False)
    db.session.commit()

    log_event("data_deletion_request", user_id=user_id, trigger="user_request",
              hidden=len(hidden_clinical), erased_sessions=len(erasable))
    for sid in hidden_clinical:
        log_event("session_hidden_by_user", session_id=sid, user_id=user_id, trigger="user_request")

    session.clear()
    return jsonify({"hidden": len(hidden_clinical), "erased": len(erasable)}), 200


# ---------------------------------------------------------------------------
# Routes — transcript download
# ---------------------------------------------------------------------------

def _transcript_data(session_id):
    """Return (messages, mode, generated_at) for a given session_id."""
    messages = (
        ChatMessage.query
        .filter_by(session_id=session_id)
        .order_by(ChatMessage.timestamp.asc())
        .all()
    )
    ts = db.session.get(TherapySession, session_id)
    # Use the session's own mode (authoritative). The creator of a therapist-led
    # session is a Clinician, not a User row, so deriving mode from User.therapy_mode
    # yielded "Unknown" for those sessions.
    mode = ts.mode.capitalize() if ts and ts.mode else "Unknown"
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    return messages, mode, generated_at


_FONT_DIR = os.path.join(os.path.dirname(__file__), "static", "fonts")
_FONT_REGULAR = os.path.join(_FONT_DIR, "DejaVuSans.ttf")
_FONT_BOLD    = os.path.join(_FONT_DIR, "DejaVuSans-Bold.ttf")


def _user_can_access_session(session_id: str, user_id: str) -> bool:
    """A user can download a session's transcript if they created it OR
    posted any message in it. Covers solo (creator-only) and couple/group
    (creator + every participant)."""
    if not user_id:
        return False
    ts = db.session.get(TherapySession, session_id)
    if not ts:
        return False
    if ts.created_by == user_id:
        return True
    # The clinician who led the session can always retrieve its full record.
    if ts.therapist_id and ts.therapist_id == user_id:
        return True
    # A participant who hid this session can no longer retrieve it (their own view
    # only — the clinician's record is untouched).
    if db.session.query(SessionHidden.id).filter_by(session_id=session_id, user_id=user_id).first():
        return False
    return ChatMessage.query.filter_by(session_id=session_id, user_id=user_id).first() is not None


# ---------------------------------------------------------------------------
# Therapist-only session summary (clinical recap + grounded ICD codes for
# billing reference + a plain-language draft the therapist MAY share). Never
# reaches the client through the system.
# ---------------------------------------------------------------------------

def _session_transcript_text(messages, ts) -> str:
    """Full speaker-labelled transcript for the summary model."""
    therapist_id = ts.therapist_id if ts else None
    lines = []
    for m in messages:
        if m.user_id == therapist_id:
            who = "Therapist"
        elif m.user_id == "AI":
            who = "AI"
        else:
            who = m.display_name or "Client"
        lines.append(f"{who}: {m.text}")
    return "\n".join(lines)


def _surfaced_codes(session_id: str) -> list:
    """Distinct ICD reference codes that actually surfaced in this session.

    Read from the persisted reference cards (grounded — the codes came from the
    curated corpus, never the model), deduped by code, in first-seen order.
    """
    import json as _json
    rows = (
        CopilotCard.query
        .filter_by(session_id=session_id, card_type="reference")
        .order_by(CopilotCard.id.asc())
        .all()
    )
    seen, out = set(), []
    for r in rows:
        code = source = ""
        try:
            payload = _json.loads(r.payload)
            code = payload.get("code", "")
            source = payload.get("source", "")
        except Exception:
            pass
        if not code or code in seen:
            continue
        seen.add(code)
        match = re.search(r"associated with (.+?)\.", r.text or "")
        out.append({"label": match.group(1) if match else "", "code": code, "source": source})
    return out


def _session_copilot_cards(session_id: str) -> list:
    """Every persisted co-pilot card for the record (safety/risk alerts, ICD
    references and AI suggestions), oldest first — so the downloaded files hold the
    full alert history even when the live panel was throttled or stopped."""
    import json as _json
    rows = (CopilotCard.query
            .filter_by(session_id=session_id)
            .order_by(CopilotCard.created_at.asc(), CopilotCard.id.asc())
            .all())
    out = []
    for r in rows:
        code = ""
        try:
            code = (_json.loads(r.payload) or {}).get("code", "")
        except Exception:
            pass
        out.append({"type": r.card_type or "observation", "text": r.text or "", "code": code})
    return out


def _session_summary_payload(session_id: str, ts) -> dict:
    """Therapist-only summary payload, cached per session and reused while the
    conversation is unchanged (keyed on message count) — generation is a slow LLM
    call, so this keeps repeat console views and downloads instant. The grounded
    codes list is always present; the AI narrative is best-effort."""
    import json as _json
    msg_count = ChatMessage.query.filter_by(session_id=session_id).count()

    cached = db.session.get(SessionSummary, session_id)
    if cached is not None and cached.message_count == msg_count:
        try:
            data = _json.loads(cached.payload)
            data["cached"] = True
            return data
        except Exception:
            pass  # corrupt cache → fall through and regenerate

    messages, _mode, generated_at = _transcript_data(session_id)
    transcript = _session_transcript_text(messages, ts)
    codes = _surfaced_codes(session_id)
    summary = clinical_summary.generate(transcript, codes, mode=ts.mode if ts else "solo") or {}
    log_event("session_summary_generated", session_id=session_id,
              user_id=ts.therapist_id if ts else None,
              codes=len(codes), narrative=bool(summary))
    payload = {
        "session_id": session_id,
        "generated_at": generated_at,
        "disclaimer": clinical_summary.DISCLAIMER,
        "codes": codes,
        "clinical": summary.get("clinical", ""),
        "codes_rationale": summary.get("codes_rationale", ""),
        "client_recap": summary.get("client_recap", ""),
        "narrative_available": bool(summary),
        "cached": False,
    }
    _store_summary(session_id, payload, msg_count)
    return payload


def _store_summary(session_id: str, payload: dict, msg_count: int) -> None:
    """Upsert the cached summary. Best-effort — a cache write must never break the
    console request or the download."""
    import json as _json
    try:
        row = db.session.get(SessionSummary, session_id)
        if row is None:
            row = SessionSummary(session_id=session_id)
            db.session.add(row)
        row.payload = _json.dumps(payload)
        row.message_count = msg_count
        row.generated_at = datetime.now(timezone.utc)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        app.logger.error("session summary cache store failed: %s", type(e).__name__)


def _is_session_therapist(session_id: str):
    """Return the TherapySession iff the current user is its therapist, else None."""
    ts = db.session.get(TherapySession, session_id)
    if ts and ts.therapist_id and ts.therapist_id == session.get("user_id"):
        return ts
    return None


@app.route("/session/<session_id>/summary")
def session_summary(session_id):
    """Therapist-only: generate and return the private session summary as JSON."""
    ts = _is_session_therapist(session_id)
    if ts is None:
        return jsonify({"error": "Forbidden"}), 403
    if not _has_ai_analysis(_session_clinician(session_id)):
        return jsonify({
            "locked": True,
            "message": "The AI session summary is part of the Pro plan.",
            "upgrade_url": url_for("billing_page"),
        }), 402
    return jsonify(_session_summary_payload(session_id, ts))


# ---------------------------------------------------------------------------
# Phase 4 — session recording (therapist-only, behind RECORDING_ENABLED flag).
# Records the live room via self-hosted LiveKit Egress → MP4 in the recordings
# bucket. Step 2 adds the all-party consent gate; Step 3 adds 30-day retention.
# ---------------------------------------------------------------------------

@app.route("/session/<session_id>/recording/start", methods=["POST"])
def recording_start(session_id):
    if not config.RECORDING_ENABLED:
        return jsonify({"error": "recording_disabled"}), 403
    if _is_session_therapist(session_id) is None:
        return jsonify({"error": "Forbidden"}), 403
    if not _has_recording(_session_clinician(session_id)):
        return jsonify({"error": "recording_requires_premium"}), 402
    # NOTE: all-party consent gating is added in Phase 4 Step 2.
    now = datetime.now(timezone.utc)
    filepath = f"{session_id}/{now.strftime('%Y%m%dT%H%M%SZ')}.mp4"
    egress_id = recording.start_recording(session_id, filepath)
    row = SessionRecording(
        session_id=session_id, egress_id=egress_id, gcs_object=filepath,
        status="active" if egress_id else "failed",
        started_by=session.get("user_id"), started_at=now,
        download_token=secrets.token_urlsafe(32),
    )
    db.session.add(row)
    db.session.commit()
    if not egress_id:
        return jsonify({"error": "recording_start_failed"}), 502
    log_event("recording_started", session_id=session_id, user_id=session.get("user_id"))
    return jsonify({"recording_id": row.id, "status": "active"}), 200


@app.route("/session/<session_id>/recording/stop", methods=["POST"])
def recording_stop(session_id):
    if not config.RECORDING_ENABLED:
        return jsonify({"error": "recording_disabled"}), 403
    if _is_session_therapist(session_id) is None:
        return jsonify({"error": "Forbidden"}), 403
    row = (
        SessionRecording.query
        .filter_by(session_id=session_id, status="active")
        .order_by(SessionRecording.id.desc())
        .first()
    )
    if row is None:
        return jsonify({"error": "no_active_recording"}), 404
    ok = recording.stop_recording(row.egress_id)
    if ok:
        row.status = "stopped"
        row.stopped_at = datetime.now(timezone.utc)
        db.session.commit()
        _finalize_stopped_recording(row)   # stamp 30-day retention + email the link
        log_event("recording_stopped", session_id=session_id, user_id=session.get("user_id"))
    return jsonify({"stopped": ok}), 200


# ---------------------------------------------------------------------------
# Recording retention (Phase 4 Step 3) — 30-day lifecycle.
#   • On stop: stamp retention_expires_at and email the clinician the link.
#   • Daily sweep: 24h-before-deletion reminder, then delete the object at expiry.
# All paths are no-ops unless RECORDING_ENABLED. None of these ever raise into
# their callers (a notification problem must never break a live session).
# ---------------------------------------------------------------------------

RECORDING_RETENTION_DAYS = 30


def _recording_download_url(row) -> str:
    """Absolute link to the in-app, therapist-gated download route. Keyed on the
    recording's opaque token (NOT the session id), so the session id never appears
    in the URL/email. Still requires the clinician to be logged in."""
    return f"{config.PUBLIC_BASE_URL}/recording/download/{row.download_token}"


def _recording_email_content(row, kind: str):
    """Return (subject, plain, html) for the 'ready' or 'reminder' email."""
    from html import escape as h
    video_url = _recording_download_url(row)
    pdf_url   = f"{config.PUBLIC_BASE_URL}/recording/download/{row.download_token}/pdf"
    docx_url  = f"{config.PUBLIC_BASE_URL}/recording/download/{row.download_token}/docx"
    expires = row.retention_expires_at
    expires_str = expires.strftime("%d %b %Y") if expires else "30 days from recording"
    if kind == "reminder":
        subject = "TogetherMindsAI — your session recording is deleted tomorrow"
        lead = ("This is a final reminder: the session recording below is scheduled to be "
                f"permanently deleted on {expires_str}. Download it now if you still need it.")
    else:
        subject = "TogetherMindsAI — your session recording is ready"
        lead = ("Your session recording is ready to download. For privacy it is kept for "
                f"{RECORDING_RETENTION_DAYS} days and then permanently deleted "
                f"(on or about {expires_str}).")
    plain = (
        f"{lead}\n\n"
        f"Session: {row.session_id}\n"
        "Sign in as this session's clinician to open these:\n"
        f"  • Video recording: {video_url}\n"
        f"  • Transcript + AI analysis + ICD codes (PDF): {pdf_url}\n"
        f"  • Transcript + AI analysis + ICD codes (Word): {docx_url}\n\n"
        "Only you, signed in as this session's clinician, can open these links.\n"
    )
    _link = "color:#2e7d32;text-decoration:none;font-weight:600;"
    html_body = (
        '<div style="font-family:-apple-system,Segoe UI,Roboto,sans-serif;max-width:640px;">'
        f'<h2 style="color:#2e7d32;">{h(subject.split("— ")[-1].capitalize())}</h2>'
        f'<p style="color:#212121;line-height:1.5;">{h(lead)}</p>'
        f'<p style="margin:18px 0 6px;"><a href="{h(video_url)}" '
        'style="background:#2e7d32;color:#fff;text-decoration:none;padding:10px 18px;'
        'border-radius:8px;display:inline-block;">▶ Download video</a></p>'
        f'<p style="margin:6px 0;"><a href="{h(pdf_url)}" style="{_link}">📄 Transcript, AI analysis &amp; ICD codes (PDF)</a></p>'
        f'<p style="margin:6px 0;"><a href="{h(docx_url)}" style="{_link}">📝 Transcript, AI analysis &amp; ICD codes (Word)</a></p>'
        f'<p style="color:#555;font-size:13px;margin-top:14px;">Session <strong>{h(row.session_id)}</strong>. '
        'Only you, signed in as this session\'s clinician, can open these links.</p>'
        '<p style="color:#888;font-size:12px;margin-top:20px;">Sent automatically by TogetherMindsAI.</p>'
        "</div>"
    )
    return subject, plain, html_body


def _email_recording(row, kind: str) -> bool:
    """Email the owning clinician the 'ready' or 'reminder' notice. Returns True
    on a successful send, False if skipped (no email/SMTP) or on failure. Assumes
    an active app context. Never raises."""
    try:
        if not (config.FEEDBACK_SMTP_USER and config.FEEDBACK_SMTP_PASSWORD):
            return False
        clinician = db.session.get(Clinician, row.started_by) if row.started_by else None
        email = getattr(clinician, "email", None)
        if not email:
            return False
        subject, plain, html_body = _recording_email_content(row, kind)
        _send_email([email], subject, plain, html_body)
        return True
    except Exception:
        app.logger.warning("recording %s email failed (id=%s)", kind, getattr(row, "id", "?"))
        return False


def _dispatch_recording_ready(rec_id: int) -> None:
    """Send the 'ready' email off the request path (SMTP can take seconds)."""
    if not config.RECORDING_ENABLED:
        return

    def _run():
        with app.app_context():
            row = db.session.get(SessionRecording, rec_id)
            if row and row.status == "stopped":
                _email_recording(row, "ready")

    threading.Thread(target=_run, daemon=True).start()


def _finalize_stopped_recording(row) -> None:
    """Stamp the 30-day retention deadline on a just-stopped recording and email
    the clinician the download link. Never raises into the caller."""
    try:
        base = row.stopped_at or datetime.now(timezone.utc)
        row.retention_expires_at = base + timedelta(days=RECORDING_RETENTION_DAYS)
        db.session.commit()
        _dispatch_recording_ready(row.id)
    except Exception:
        db.session.rollback()
        app.logger.warning("finalize recording failed (id=%s)", getattr(row, "id", "?"))


def _recording_retention_sweep() -> None:
    """Daily: send due 24h-before-deletion reminders, then delete recordings past
    their retention deadline. Idempotent and safe to run repeatedly. Never raises."""
    if not config.RECORDING_ENABLED:
        return
    try:
        with app.app_context():
            now = datetime.now(timezone.utc)
            soon = now + timedelta(hours=24)

            # 1) Reminders — expiring within 24h, not yet reminded, not deleted.
            due = (SessionRecording.query
                   .filter(SessionRecording.status == "stopped",
                           SessionRecording.retention_expires_at.isnot(None),
                           SessionRecording.retention_expires_at <= soon,
                           SessionRecording.retention_expires_at > now,
                           SessionRecording.reminder_sent_at.is_(None))
                   .all())
            for row in due:
                if _email_recording(row, "reminder"):
                    row.reminder_sent_at = now
                    db.session.commit()

            # 2) Deletions — past the retention deadline.
            expired = (SessionRecording.query
                       .filter(SessionRecording.status == "stopped",
                               SessionRecording.retention_expires_at.isnot(None),
                               SessionRecording.retention_expires_at <= now)
                       .all())
            for row in expired:
                if recording.delete_object(row.gcs_object):
                    row.status = "deleted"
                    db.session.commit()
                    log_event("recording_deleted", session_id=row.session_id,
                              user_id=row.started_by, recording_id=row.id,
                              trigger="retention_expired")
    except Exception as exc:
        db.session.rollback()
        app.logger.error("recording retention sweep error: %s", type(exc).__name__)


@app.route("/recording/download/<token>")
def recording_download(token):
    """Therapist-only streamed download of a session recording, keyed on an opaque
    token so the session id never appears in the URL. The object is streamed
    through the app so the existing login/authorization applies (no raw signed URL
    is ever handed out)."""
    if not config.RECORDING_ENABLED:
        abort(404)
    row = SessionRecording.query.filter_by(download_token=token).first() if token else None
    if row is None:
        abort(404)
    # Still therapist-gated: the logged-in user must be the session's clinician.
    if _is_session_therapist(row.session_id) is None:
        abort(403)
    if row.status == "deleted" or not row.gcs_object:
        abort(410)   # Gone — past its 30-day retention
    gen, size, ctype = recording.download_stream(row.gcs_object)
    if gen is None:
        abort(502)
    log_event("recording_downloaded", session_id=row.session_id,
              user_id=session.get("user_id"), recording_id=row.id)
    resp = Response(stream_with_context(gen), mimetype=ctype or "video/mp4")
    resp.headers["Content-Disposition"] = f'attachment; filename="recording-{row.id}.mp4"'
    if size:
        resp.headers["Content-Length"] = str(size)
    return resp


# Wire the daily recording-retention sweep now that its function is defined (the
# `_scheduler` global was created in the startup block above). Registered here,
# not in that block, because the sweep is defined further down this module.
if not config.IS_TESTING:
    _scheduler.add_job(
        _recording_retention_sweep, "interval", hours=24, id="recording_retention_sweep",
    )
    # Catch up on any reminders/deletions that came due while the app was down.
    threading.Thread(target=_recording_retention_sweep, daemon=True).start()


def _render_summary_pdf(pdf, summary: dict) -> None:
    """Prepend the therapist-only summary to the PDF, above the transcript."""
    from fpdf.enums import XPos, YPos

    # multi_cell defaults to new_x=RIGHT, which leaves the cursor at the right
    # margin and makes the NEXT multi_cell raise "not enough horizontal space".
    # Return to the left margin after every block.
    mc = {"new_x": XPos.LMARGIN, "new_y": YPos.NEXT}

    pdf.set_font("DejaVu", "B", 14)
    pdf.set_text_color(146, 39, 15)
    pdf.cell(0, 9, "Clinician Summary — Private (therapist only)",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("DejaVu", "", 9)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(0, 5, summary.get("disclaimer", ""), **mc)
    pdf.ln(2)

    def _section(title, body):
        if not body:
            return
        pdf.set_font("DejaVu", "B", 11)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(0, 6, title, **mc)
        pdf.set_font("DejaVu", "", 10)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(0, 5, body, **mc)
        pdf.ln(2)

    _section("Clinical summary",
             summary.get("clinical") or "(AI narrative unavailable — see transcript below.)")

    # ICD codes — always rendered (grounded data), even if the narrative failed.
    pdf.set_font("DejaVu", "B", 11)
    pdf.set_text_color(40, 40, 40)
    pdf.multi_cell(0, 6, "ICD codes (billing reference)", **mc)
    pdf.set_font("DejaVu", "", 10)
    pdf.set_text_color(30, 30, 30)
    codes = summary.get("codes") or []
    if codes:
        for c in codes:
            label = f"{c['label']} — " if c.get("label") else ""
            pdf.multi_cell(0, 5, f"• {label}{c.get('code', '')}  ({c.get('source', '')})", **mc)
    else:
        pdf.multi_cell(0, 5, "No ICD reference codes surfaced during this session.", **mc)
    if summary.get("codes_rationale"):
        pdf.ln(1)
        pdf.multi_cell(0, 5, summary["codes_rationale"], **mc)
    pdf.ln(2)

    # Co-pilot alerts — the full saved record (incl. any not shown live).
    cards = summary.get("copilot_cards") or []
    if cards:
        _CARD_LABEL = {"risk": "Risk", "reference": "Reference",
                       "suggestion": "Suggestion", "observation": "Observation"}
        pdf.set_font("DejaVu", "B", 11)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(0, 6, "Co-pilot alerts (full record)", **mc)
        pdf.set_font("DejaVu", "", 10)
        pdf.set_text_color(30, 30, 30)
        for c in cards:
            lbl = _CARD_LABEL.get(c.get("type"), (c.get("type") or "Note").title())
            code = f"  [{c['code']}]" if c.get("code") else ""
            pdf.multi_cell(0, 5, f"• [{lbl}] {c.get('text', '')}{code}", **mc)
        pdf.ln(2)

    _section("Client-facing draft — share only at your discretion "
             "(the client cannot see this unless you give it to them)",
             summary.get("client_recap"))

    pdf.set_draw_color(200, 200, 200)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)
    pdf.set_font("DejaVu", "B", 13)
    pdf.set_text_color(30, 30, 30)
    pdf.cell(0, 8, "Transcript", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)


def _render_summary_docx(doc, summary: dict) -> None:
    """Prepend the therapist-only summary to the DOCX, above the transcript."""
    from docx.shared import Pt, RGBColor

    h = doc.add_heading("Clinician Summary — Private (therapist only)", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x92, 0x27, 0x0F)
    disc = doc.add_paragraph()
    run = disc.add_run(summary.get("disclaimer", ""))
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x78, 0x78, 0x78)

    def _section(title, body):
        if not body:
            return
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(body)

    _section("Clinical summary",
             summary.get("clinical") or "(AI narrative unavailable — see transcript below.)")

    p = doc.add_paragraph()
    p.add_run("ICD codes (billing reference)").bold = True
    codes = summary.get("codes") or []
    if codes:
        for c in codes:
            label = f"{c['label']} — " if c.get("label") else ""
            doc.add_paragraph(f"• {label}{c.get('code', '')}  ({c.get('source', '')})")
    else:
        doc.add_paragraph("No ICD reference codes surfaced during this session.")
    if summary.get("codes_rationale"):
        doc.add_paragraph(summary["codes_rationale"])

    # Co-pilot alerts — the full saved record (incl. any not shown live).
    cards = summary.get("copilot_cards") or []
    if cards:
        _CARD_LABEL = {"risk": "Risk", "reference": "Reference",
                       "suggestion": "Suggestion", "observation": "Observation"}
        p = doc.add_paragraph()
        p.add_run("Co-pilot alerts (full record)").bold = True
        for c in cards:
            lbl = _CARD_LABEL.get(c.get("type"), (c.get("type") or "Note").title())
            code = f"  [{c['code']}]" if c.get("code") else ""
            doc.add_paragraph(f"• [{lbl}] {c.get('text', '')}{code}")

    _section("Client-facing draft — share only at your discretion "
             "(the client cannot see this unless you give it to them)",
             summary.get("client_recap"))

    doc.add_paragraph("─" * 40)
    doc.add_heading("Transcript", level=2)


def _transcript_pdf_buf(session_id: str) -> io.BytesIO:
    """Render the session transcript as a PDF in memory. Includes the therapist's
    private summary + co-pilot alert record when the viewer is the clinician."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    messages, mode, generated_at = _transcript_data(session_id)

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_font("DejaVu",      fname=_FONT_REGULAR)
    pdf.add_font("DejaVu", "B", fname=_FONT_BOLD)
    pdf.add_page()

    # Title
    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 10, "TogetherMindsAI \u2014 Session Transcript",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    # Metadata
    pdf.set_font("DejaVu", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Session ID : {session_id}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 6, f"Mode       : {mode}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 6, f"Generated  : {generated_at}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    # Divider
    pdf.set_draw_color(200, 200, 200)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)

    # Therapist-only: prepend the private clinical summary + grounded ICD codes.
    # The AI summary is a paid (Pro/Premium) feature; the transcript itself is free.
    _ts_therapist = _is_session_therapist(session_id)
    if _ts_therapist is not None and _has_ai_analysis(_session_clinician(session_id)):
        _sum = _session_summary_payload(session_id, _ts_therapist)
        _sum["copilot_cards"] = _session_copilot_cards(session_id)   # full alert record
        _render_summary_pdf(pdf, _sum)

    # Assign a distinct RGB color to each human participant
    _PDF_PARTICIPANT_COLORS = [
        (30,  80,  160),   # blue
        (146, 39,  15),    # rust red
        (107, 33,  168),   # purple
        (13,  110, 110),   # teal
        (180, 90,  9),     # amber
        (26,  92,  46),    # dark green
    ]
    pdf_participant_color: dict = {}
    _pdf_color_idx = 0
    for msg in messages:
        if msg.user_id != "AI" and msg.user_id not in pdf_participant_color:
            pdf_participant_color[msg.user_id] = _PDF_PARTICIPANT_COLORS[_pdf_color_idx % len(_PDF_PARTICIPANT_COLORS)]
            _pdf_color_idx += 1

    if not messages:
        pdf.set_text_color(120, 120, 120)
        pdf.set_font("DejaVu", "", 11)
        pdf.cell(0, 8, "No messages recorded for this session.",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    else:
        for msg in messages:
            is_ai = msg.user_id == "AI"
            if is_ai:
                speaker = "AI Co-Pilot"
            elif msg.display_name:
                speaker = f"{session_id}-{msg.display_name}"
            else:
                speaker = "User"
            ts = msg.timestamp.strftime("%Y-%m-%d %H:%M")

            # Speaker + timestamp line
            pdf.set_font("DejaVu", "B", 10)
            if is_ai:
                pdf.set_text_color(30, 120, 60)
            else:
                r, g, b = pdf_participant_color.get(msg.user_id, (30, 80, 160))
                pdf.set_text_color(r, g, b)
            pdf.cell(0, 7, f"{speaker}  [{ts}]",
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            # Message body
            pdf.set_font("DejaVu", "", 11)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 6, msg.text)
            pdf.ln(3)

    buf = io.BytesIO(pdf.output())
    buf.seek(0)
    return buf


@app.route("/transcript/<session_id>/pdf")
def download_transcript_pdf(session_id):
    if not _user_can_access_session(session_id, session.get("user_id")):
        return jsonify({"error": "Forbidden"}), 403
    # PHI disclosure — record it (HIPAA § 164.312(b) access logging).
    log_event("transcript_downloaded", session_id=session_id,
              user_id=session.get("user_id"), format="pdf")
    filename = f"transcript_{session_id}_{datetime.now(timezone.utc).strftime('%Y%m%d')}.pdf"
    # send_file streams via the WSGI file wrapper and supports range/conditional
    # requests, so the browser's ranged download completes instead of resetting.
    return send_file(_transcript_pdf_buf(session_id), mimetype="application/pdf",
                     as_attachment=True, download_name=filename)


def _transcript_docx_buf(session_id: str) -> io.BytesIO:
    """Render the session transcript as a DOCX in memory. Includes the therapist's
    private summary + co-pilot alert record when the viewer is the clinician."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    messages, mode, generated_at = _transcript_data(session_id)

    doc = Document()

    # Title
    title = doc.add_heading("TogetherMindsAI — Session Transcript", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x78, 0x40)

    # Metadata table
    meta = doc.add_paragraph()
    meta.add_run("Session ID : ").bold = True
    meta.add_run(session_id)
    meta2 = doc.add_paragraph()
    meta2.add_run("Mode       : ").bold = True
    meta2.add_run(mode)
    meta3 = doc.add_paragraph()
    meta3.add_run("Generated  : ").bold = True
    meta3.add_run(generated_at)

    doc.add_paragraph("─" * 40)

    # Therapist-only: prepend the private clinical summary + grounded ICD codes.
    # The AI summary is a paid (Pro/Premium) feature; the transcript itself is free.
    _ts_therapist = _is_session_therapist(session_id)
    if _ts_therapist is not None and _has_ai_analysis(_session_clinician(session_id)):
        _sum = _session_summary_payload(session_id, _ts_therapist)
        _sum["copilot_cards"] = _session_copilot_cards(session_id)   # full alert record
        _render_summary_docx(doc, _sum)

    # Assign a distinct color to each human participant (AI is always green)
    _PARTICIPANT_COLORS = [
        RGBColor(0x1E, 0x50, 0xA0),  # blue
        RGBColor(0x92, 0x27, 0x0F),  # rust red
        RGBColor(0x6B, 0x21, 0xA8),  # purple
        RGBColor(0x0D, 0x6E, 0x6E),  # teal
        RGBColor(0xB4, 0x5A, 0x09),  # amber
        RGBColor(0x1A, 0x5C, 0x2E),  # dark green
    ]
    participant_color: dict = {}
    _color_idx = 0
    for msg in messages:
        if msg.user_id != "AI" and msg.user_id not in participant_color:
            participant_color[msg.user_id] = _PARTICIPANT_COLORS[_color_idx % len(_PARTICIPANT_COLORS)]
            _color_idx += 1

    if not messages:
        p = doc.add_paragraph("No messages recorded for this session.")
        p.runs[0].italic = True
    else:
        for msg in messages:
            is_ai = msg.user_id == "AI"
            if is_ai:
                speaker = "AI Co-Pilot"
            elif msg.display_name:
                speaker = f"{session_id}-{msg.display_name}"
            else:
                speaker = "User"
            ts = msg.timestamp.strftime("%Y-%m-%d %H:%M")

            # Speaker heading
            p = doc.add_paragraph()
            run = p.add_run(f"{speaker}  [{ts}]")
            run.bold = True
            run.font.size = Pt(10)
            if is_ai:
                run.font.color.rgb = RGBColor(0x1E, 0x78, 0x3C)
            else:
                run.font.color.rgb = participant_color.get(msg.user_id, RGBColor(0x1E, 0x50, 0xA0))

            # Message body
            doc.add_paragraph(msg.text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/transcript/<session_id>/docx")
def download_transcript_docx(session_id):
    if not _user_can_access_session(session_id, session.get("user_id")):
        return jsonify({"error": "Forbidden"}), 403
    # PHI disclosure — record it (HIPAA § 164.312(b) access logging).
    log_event("transcript_downloaded", session_id=session_id,
              user_id=session.get("user_id"), format="docx")
    filename = f"transcript_{session_id}_{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"
    return send_file(
        _transcript_docx_buf(session_id),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True, download_name=filename)


@app.route("/recording/download/<token>/<fmt>")
def recording_download_doc(token, fmt):
    """Token-keyed transcript download (PDF/Word) tied to a recording, so the
    recording email can link the documents without exposing the session id.
    Therapist-gated, same as the video download."""
    if fmt not in ("pdf", "docx"):
        abort(404)
    if not config.RECORDING_ENABLED:
        abort(404)
    row = SessionRecording.query.filter_by(download_token=token).first() if token else None
    if row is None:
        abort(404)
    if _is_session_therapist(row.session_id) is None:
        abort(403)
    log_event("transcript_downloaded", session_id=row.session_id,
              user_id=session.get("user_id"), format=fmt, via="recording_token")
    if fmt == "pdf":
        return send_file(_transcript_pdf_buf(row.session_id), mimetype="application/pdf",
                         as_attachment=True, download_name=f"session-{row.id}.pdf")
    return send_file(
        _transcript_docx_buf(row.session_id),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True, download_name=f"session-{row.id}.docx")


# ---------------------------------------------------------------------------
# Routes — public-key authentication API
# ---------------------------------------------------------------------------

@app.route("/api/auth/register", methods=["POST"])
@limiter.limit("10 per hour")
def api_auth_register():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Invalid request"}), 400

    public_key_b64 = (data.get("public_key") or "").strip()
    therapy_mode   = (data.get("therapy_mode") or "solo").strip()

    if not public_key_b64:
        return jsonify({"error": "public_key required"}), 400
    if therapy_mode not in ("solo", "couple", "group"):
        return jsonify({"error": "Invalid therapy_mode"}), 400
    try:
        base64.b64decode(public_key_b64, validate=True)
    except Exception:
        return jsonify({"error": "Invalid public_key encoding"}), 400

    # Registration only JOINS a clinician-led session that /session/join stashed as
    # pending_*. It never creates a new self-directed session (the sole creator is a
    # logged-in clinician via POST /therapist/start/<mode>).
    pending_couple = session.pop("pending_couple_session", None)
    pending_group  = session.pop("pending_group_session",  None)
    pending_solo   = session.pop("pending_solo_session",   None)
    joined_sid = pending_solo or pending_couple or pending_group
    if not joined_sid:
        return jsonify({"error": "No session to join"}), 400

    user_id = str(uuid.uuid4())
    db.session.add(User(id=user_id, therapy_mode=therapy_mode, public_key=public_key_b64))
    db.session.commit()
    session["user_id"] = user_id

    return jsonify({
        "user_id": user_id,
        "therapy_mode": therapy_mode,
        "session_id": joined_sid,
    }), 201


@app.route("/api/auth/challenge", methods=["POST"])
def api_auth_challenge():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Invalid request"}), 400

    user_id = (data.get("user_id") or "").strip()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404
    if not user.public_key:
        return jsonify({"error": "No public key registered for this user"}), 400

    nonce = secrets.token_hex(32)
    user.challenge = nonce
    user.challenge_expires_at = time.time() + 300   # 5-minute window
    db.session.commit()

    return jsonify({"challenge": nonce}), 200


@app.route("/api/auth/verify", methods=["POST"])
def api_auth_verify():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Invalid request"}), 400

    user_id       = (data.get("user_id") or "").strip()
    signature_b64 = (data.get("signature") or "").strip()

    user = db.session.get(User, user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404
    if not user.challenge or not user.challenge_expires_at:
        return jsonify({"error": "No active challenge"}), 401
    if time.time() > user.challenge_expires_at:
        user.challenge = None
        user.challenge_expires_at = None
        db.session.commit()
        return jsonify({"error": "Challenge expired"}), 401

    try:
        pub_key_bytes = base64.b64decode(user.public_key)
        public_key    = load_der_public_key(pub_key_bytes)
        sig_bytes     = base64.b64decode(signature_b64)
        public_key.verify(sig_bytes, user.challenge.encode("utf-8"), ECDSA(SHA256()))
    except (InvalidSignature, Exception):
        return jsonify({"error": "Invalid signature"}), 401

    user.challenge = None
    user.challenge_expires_at = None
    db.session.commit()

    session["user_id"] = user_id

    ts = TherapySession.query.filter_by(created_by=user_id, mode=user.therapy_mode).first()
    session_id = ts.id if ts else None

    return jsonify({
        "ok": True,
        "user_id": user_id,
        "therapy_mode": user.therapy_mode,
        "session_id": session_id,
    }), 200


# ---------------------------------------------------------------------------
# SocketIO Events
# ---------------------------------------------------------------------------

@socketio.on("join")
def on_join(data):
    try:
        session_id = data.get("session_id")
        user_id    = data.get("user_id")
        mode       = data.get("mode", "solo")

        ts = db.session.get(TherapySession, session_id)
        therapist_id = ts.therapist_id if ts else None
        eff_mode = ts.mode if (ts and ts.mode) else mode

        # Hard capacity gate: reject a client that would exceed the mode's cap
        # (solo=1, couple=2, group=unlimited). The therapist is never blocked.
        if _session_is_full(session_id, eff_mode, user_id, therapist_id):
            cap = _MODE_CLIENT_CAP.get(eff_mode)
            emit("error", {"message":
                 "This session is full — it allows up to %d participant%s."
                 % (cap, "" if cap == 1 else "s")})
            return

        join_room(session_id)
        room_mode[session_id] = mode

        # Track presence
        room_participants[session_id].add(user_id)
        sid_to_user[request.sid]    = user_id
        sid_to_session[request.sid] = session_id

        # Durable participation link (so "my sessions" finds it even if the
        # participant never sends a message).
        _record_participation(session_id, user_id)

        # Therapist co-pilot: if this session is therapist-led, remember who the
        # therapist is and — when the joiner IS the therapist — add them to the
        # private console room. Clients are never added to this room, so the
        # suggestion cards emitted there can never reach them. (ts loaded above.)
        is_therapist_led = bool(ts and ts.therapist_id)
        if is_therapist_led:
            session_therapist_id[session_id] = ts.therapist_id
            if user_id == ts.therapist_id:
                join_room(_therapist_room(session_id))
                emit("console_init", {"mode": mode})
                # Replay any stored cards so the console shows full history, not
                # just whatever arrived live since this connection opened.
                _emit_card_history(session_id)

        # Assign join position (used for default display names)
        joined = session_joined_users.setdefault(session_id, [])
        if user_id not in joined:
            joined.append(user_id)
        join_position = joined.index(user_id) + 1
        # The clinician leading the session defaults to "Therapist" (in all modes)
        # so their role is clear to everyone; clients keep the mode-based name.
        if therapist_id and user_id == therapist_id:
            default_name = "Therapist"
        else:
            default_name = _default_display_name(mode, join_position)

        messages = (
            ChatMessage.query
            .filter_by(session_id=session_id)
            .order_by(ChatMessage.timestamp.asc())
            .all()
        )
        current_display_name = session_display_names.get(session_id, {}).get(user_id)
        emit("history", {
            "messages": [m.to_dict() for m in messages],
            "join_position": join_position,
            "default_name": default_name,
            "current_display_name": current_display_name,
        })
        emit("participant_list",
             {"participants": list(room_participants[session_id])},
             to=session_id)
        emit("participant_joined", {"user_id": user_id}, to=session_id)

        # If a recording is in progress/requested, a newcomer must consent too —
        # prompt them, and re-evaluate (an unconsented newcomer pauses recording).
        if session_recording_requested.get(session_id):
            emit("recording_consent_prompt", {"requested_by": "Therapist"})
            _evaluate_recording(session_id)
        _emit_recording_state(session_id)

        # Sync the shared session friendly name to this newcomer (silent — no popup).
        _fn = session_friendly_name.get(session_id)
        if _fn:
            emit("friendly_name_set", {"name": _fn, "silent": True})
    except Exception as e:
        app.logger.error("on_join error: %s", type(e).__name__)
        emit("error", {"message": "Failed to join session. Please refresh."})


@socketio.on("disconnect")
def on_disconnect():
    user_id    = sid_to_user.pop(request.sid, None)
    session_id = sid_to_session.pop(request.sid, None)
    if user_id and session_id:
        room_participants[session_id].discard(user_id)
        # session_display_names is intentionally NOT cleared on disconnect.
        # Socket.IO long-polling causes frequent disconnect/reconnect cycles;
        # clearing the name here causes the display name modal to re-appear
        # on every reconnect. session_taken_names permanently blocks reuse,
        # so keeping session_display_names across reconnects is safe.
        emit("participant_left",
             {"user_id": user_id},
             to=session_id)
        emit("participant_list",
             {"participants": list(room_participants[session_id])},
             to=session_id)
        # Drop their consent and re-evaluate: if the person who was blocking
        # recording just left, the remaining all-consenting participants resume.
        session_recording_consent.get(session_id, {}).pop(user_id, None)
        _evaluate_recording(session_id)
        _emit_recording_state(session_id)


@socketio.on("send_message")
def on_send_message(data):
    try:
        session_id = data.get("session_id")
        user_id    = data.get("user_id")
        text       = data.get("text", "").strip()

        if not text:
            return
        if len(text) > _MAX_MSG_LEN:
            emit("error", {"message": (
                f"Your message is too long ({len(text):,} characters). "
                f"Please keep it under {_MAX_MSG_LEN:,} characters."
            )})
            return
        if not _check_rate_limit(user_id):
            emit("rate_limited", {"message": "You're sending messages too quickly. Please slow down."})
            return

        now  = datetime.now(timezone.utc)
        mode = room_mode.get(session_id, "solo")

        # ---- Therapist-led sessions: the AI never replies to the room. It acts
        # as a private co-pilot, emitting suggestion cards to the therapist only.
        therapist_id = session_therapist_id.get(session_id)
        if therapist_id is not None:
            display_name = session_display_names.get(session_id, {}).get(user_id)
            user_msg = ChatMessage(
                session_id=session_id, user_id=user_id, text=text,
                timestamp=now, display_name=display_name,
            )
            db.session.add(user_msg)
            db.session.add(Exercise(
                user_id=user_id, type="realtime_chat", mode=mode, timestamp=now,
            ))
            db.session.commit()

            # The conversation itself is shared with everyone in the room.
            emit("new_message",
                 {"user_id": user_id, "text": text, "display_name": display_name,
                  "timestamp": now.strftime("%Y-%m-%d %H:%M:%S")},
                 to=session_id)
            log_event("message_sent", session_id=session_id, user_id=user_id,
                      mode=mode, message_length=len(text))

            # Every utterance drives the co-pilot — including the therapist's own,
            # so they get feedback on their interventions, not just the client's words.
            if user_id != therapist_id:
                # CLIENT spoke. Crisis safety net retained: the resources message is
                # still shown to the client (chosen "Both"), and a risk card alerts
                # the therapist. trigger_text drives the keyword risk check.
                if detect_crisis(text):
                    crisis_now = datetime.now(timezone.utc)
                    db.session.add(ChatMessage(
                        session_id=session_id, user_id="AI", text=CRISIS_RESPONSE,
                        timestamp=crisis_now,
                    ))
                    db.session.commit()
                    emit("new_message",
                         {"user_id": "AI", "text": CRISIS_RESPONSE, "display_name": None,
                          "timestamp": crisis_now.strftime("%Y-%m-%d %H:%M:%S")},
                         to=session_id)
                    log_event("crisis_detected", session_id=session_id, user_id=user_id,
                              layer="keyword", recipient="client_and_therapist")
                _run_copilot(session_id, mode, trigger_text=text, trigger_user_id=user_id)
            else:
                # THERAPIST spoke. Reflect on the intervention too. No client crisis
                # net and no keyword risk card from the therapist's own words.
                _run_copilot(session_id, mode, trigger_text=None, trigger_user_id=user_id)
            return
    except Exception as e:
        app.logger.error("on_send_message error: %s", type(e).__name__)
        db.session.rollback()
        emit("error", {"message": "Failed to send message. Please try again."})


@socketio.on("therapist_note")
def on_therapist_note(data):
    """Private note from the therapist that steers the co-pilot.

    The note is NOT broadcast to clients. It is appended to the co-pilot context
    and triggers a fresh round of suggestion cards. Rejected unless the sender is
    the session's therapist.
    """
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    text       = (data.get("text") or "").strip()
    if not text:
        return
    therapist_id = session_therapist_id.get(session_id)
    if therapist_id is None or user_id != therapist_id:
        return   # only the session therapist may add notes
    if len(text) > _MAX_MSG_LEN:
        return

    notes = session_therapist_notes.setdefault(session_id, [])
    notes.append(text)
    del notes[:-20]   # keep the most recent notes only
    _run_copilot(session_id, room_mode.get(session_id, "solo"), trigger_text=None,
                 trigger_user_id=user_id)


@socketio.on("copilot_cadence")
def on_copilot_cadence(data):
    """Therapist sets how chatty the co-pilot is LIVE: more / less / stop. Alerts
    are still always saved to the record (see _run_copilot). Therapist-only."""
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    mode       = (data.get("mode") or "").lower()
    if mode not in ("more", "less", "stop"):
        return
    if session_therapist_id.get(session_id) != user_id:
        return   # only the session therapist may change the cadence
    session_copilot_cadence[session_id] = mode
    session_copilot_emit_counter[session_id] = 0
    log_event("copilot_cadence_set", session_id=session_id, user_id=user_id, mode=mode)
    emit("copilot_cadence_set", {"mode": mode})   # back to the therapist console


@socketio.on("set_display_name")
def on_set_display_name(data):
    """First-time display name assignment for couple/group participants.

    The client emits this after the name prompt modal is confirmed.
    On success emits name_set back to the caller.
    On conflict emits name_error back to the caller only.
    """
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    name       = data.get("display_name", "").strip()

    if not name or len(name) > 40:
        emit("name_error", {"message": "Display name must be between 1 and 40 characters."})
        return

    if _is_name_taken(session_id, user_id, name):
        emit("name_error", {"message": f"'{name}' is already taken in this session. Please choose another."})
        return

    _claim_display_name(session_id, user_id, name)
    emit("name_set", {"user_id": user_id, "display_name": name})
    # Inform other participants
    emit("participant_named", {"user_id": user_id, "display_name": name}, to=session_id)


@socketio.on("rename")
def on_rename(data):
    """Rename a display name mid-session.

    On success emits name_changed to the entire room (so all bubbles update).
    On conflict emits name_error back to the caller only.
    """
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    new_name   = data.get("new_name", "").strip()

    if not new_name or len(new_name) > 40:
        emit("name_error", {"message": "Display name must be between 1 and 40 characters."})
        return

    old_name = session_display_names.get(session_id, {}).get(user_id, "")

    # No-op if same name (case-insensitive)
    if new_name.lower() == old_name.lower():
        emit("name_set", {"user_id": user_id, "display_name": new_name})
        return

    if _is_name_taken(session_id, user_id, new_name):
        emit("name_error", {"message": f"'{new_name}' is already taken in this session. Please choose another."})
        return

    _claim_display_name(session_id, user_id, new_name)
    # old_name stays in taken_names — it can never be reclaimed
    emit("name_changed", {"user_id": user_id, "old_name": old_name, "new_name": new_name}, to=session_id)


# ---------------------------------------------------------------------------
# Phase 4 — recording consent (all-party, reactive). Recording runs ONLY while
# every current participant consents; it stops the instant anyone declines,
# withdraws, or an unconsented participant is present — and resumes when all
# consent again. The clinician initiates; the AI/egress never records silently.
# ---------------------------------------------------------------------------

def _recording_state(session_id: str) -> dict:
    participants = list(room_participants.get(session_id, set()))
    consent = session_recording_consent.get(session_id, {})
    return {
        "requested": bool(session_recording_requested.get(session_id)),
        "active": bool(session_recording_active.get(session_id)),
        "awaiting": [u for u in participants if consent.get(u) is not True],
    }


def _emit_recording_state(session_id: str) -> None:
    socketio.emit("recording_state", _recording_state(session_id), to=session_id)


def _evaluate_recording(session_id: str) -> None:
    """Start/stop egress to match the live consent state. Records only while every
    current participant consents; never raises into the caller."""
    try:
        requested = bool(session_recording_requested.get(session_id))
        participants = set(room_participants.get(session_id, set()))
        consent = session_recording_consent.get(session_id, {})
        all_consent = bool(participants) and all(consent.get(u) is True for u in participants)
        active_id = session_recording_active.get(session_id)
        # Entitlement is the ultimate guard: recording never starts unless the
        # session's clinician holds the Premium plan (a no-op while billing is off).
        entitled = _has_recording(_session_clinician(session_id))

        if requested and all_consent and entitled and not active_id:
            session_recording_active[session_id] = -1   # pending — blocks re-entry
            now = datetime.now(timezone.utc)
            filepath = f"{session_id}/{now.strftime('%Y%m%dT%H%M%SZ')}.mp4"
            egress_id = recording.start_recording(session_id, filepath)
            row = SessionRecording(
                session_id=session_id, egress_id=egress_id, gcs_object=filepath,
                status="active" if egress_id else "failed",
                started_by=session_therapist_id.get(session_id), started_at=now,
                download_token=secrets.token_urlsafe(32),
            )
            db.session.add(row)
            db.session.commit()
            session_recording_active[session_id] = row.id if egress_id else None
            if egress_id:
                log_event("recording_started", session_id=session_id,
                          user_id=session_therapist_id.get(session_id), trigger="all_consented")
        elif active_id and active_id != -1 and (not requested or not all_consent):
            row = db.session.get(SessionRecording, active_id)
            if row and row.egress_id:
                recording.stop_recording(row.egress_id)
                row.status = "stopped"
                row.stopped_at = datetime.now(timezone.utc)
                db.session.commit()
                _finalize_stopped_recording(row)   # stamp retention + email the link
            session_recording_active[session_id] = None
            log_event("recording_stopped", session_id=session_id,
                      user_id=session_therapist_id.get(session_id),
                      trigger="stopped_or_consent_withdrawn")
    except Exception as e:
        db.session.rollback()
        app.logger.error("recording evaluate error: %s", type(e).__name__)


@socketio.on("recording_request")
def on_recording_request(data):
    """Clinician asks to record — everyone is then prompted to consent."""
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    if not config.RECORDING_ENABLED:
        emit("recording_unavailable", {"message": "Recording is not available for this service."})
        return
    if session_therapist_id.get(session_id) != user_id:
        return   # only the session's clinician may start recording
    if not _has_recording(_session_clinician(session_id)):
        emit("recording_unavailable", {"message": "Recording is a Premium-plan feature. Upgrade in Plans & billing."})
        return
    session_recording_requested[session_id] = True
    session_recording_consent[session_id][user_id] = True   # the clinician consents by requesting
    emit("recording_consent_prompt", {"requested_by": "Therapist"}, to=session_id)
    _evaluate_recording(session_id)
    _emit_recording_state(session_id)


@socketio.on("recording_consent")
def on_recording_consent(data):
    """A participant grants or withdraws consent (can change at any time)."""
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    if user_id not in room_participants.get(session_id, set()):
        return
    session_recording_consent[session_id][user_id] = bool(data.get("consent"))
    _evaluate_recording(session_id)
    _emit_recording_state(session_id)


@socketio.on("recording_cancel")
def on_recording_cancel(data):
    """Clinician turns recording off entirely (stops it and clears the request)."""
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    if session_therapist_id.get(session_id) != user_id:
        return
    session_recording_requested[session_id] = False
    _evaluate_recording(session_id)
    _emit_recording_state(session_id)


@socketio.on("end_session")
def on_end_session(data):
    """Only the session's clinician may end it. Notifies everyone else so their
    client shows a 'session ended' popup and returns them out."""
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    if session_therapist_id.get(session_id) != user_id:
        return   # only the clinician can end the session
    log_event("session_ended", session_id=session_id, user_id=user_id, trigger="therapist")
    emit("session_ended", {"by": "Therapist"}, to=session_id, include_self=False)


@socketio.on("set_friendly_name")
def on_set_friendly_name(data):
    """Clinician sets a shared, session-wide friendly name. Broadcast to everyone
    (clients get a popup); also stored so late joiners are synced on join."""
    session_id = data.get("session_id", "")
    user_id    = data.get("user_id", "")
    if session_therapist_id.get(session_id) != user_id:
        return   # only the clinician may name the session
    name = (data.get("name") or "").strip()[:60]
    session_friendly_name[session_id] = name
    log_event("friendly_name_set", session_id=session_id, user_id=user_id)
    emit("friendly_name_set", {"name": name, "by": "Therapist"}, to=session_id, include_self=False)


if __name__ == "__main__":
    _debug = os.environ.get("FLASK_DEBUG", "false").lower() == "true"

    _cert = os.path.join(os.path.dirname(__file__), "certs", "cert.pem")
    _key  = os.path.join(os.path.dirname(__file__), "certs", "key.pem")
    _ssl  = (_cert, _key) if os.path.exists(_cert) and os.path.exists(_key) else None
    if _ssl:
        app.logger.info("TLS enabled — https://localhost:5001  |  https://192.168.1.88:5001")
    else:
        app.logger.warning("certs/cert.pem or key.pem not found — running without TLS")

    socketio.run(app, debug=_debug, use_reloader=False, host="0.0.0.0", port=5001,
                 ssl_context=_ssl)
