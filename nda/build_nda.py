"""Generate a password-protected one-way NDA as PDF and DOCX.

The open password is read from NDA_SECRET (environment, or the repo .env) — it is
never hardcoded. Recipients need that password to open either file.

Setup:   python -m venv .venv && .venv/Scripts/python -m pip install -r requirements.txt
Run:     NDA_SECRET must be set (in the repo .env or the environment), then:
         .venv/Scripts/python build_nda.py
Output:  ./NDA.pdf  and  ./NDA.docx   (both AES-encrypted with NDA_SECRET)
"""
import io
import os
import sys

from fpdf import FPDF
from fpdf.enums import AccessPermission
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from msoffcrypto.format.ooxml import OOXMLFile

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
PDF_OUT = os.path.join(HERE, "NDA.pdf")
DOCX_OUT = os.path.join(HERE, "NDA.docx")

DISCLOSER = "TogetherMindsAI, ai4org, and Global Business Consulting, Inc."
TITLE = "NON-DISCLOSURE AGREEMENT"
SUBTITLE = "(One-Way / Unilateral)"

INTRO = (
    f"This Non-Disclosure Agreement (“Agreement”) is entered into as of "
    f"[Effective Date] by and between {DISCLOSER} (collectively, the “Disclosing "
    f"Party”), and the undersigned recipient (“Receiving Party”). The "
    f"Disclosing Party and the Receiving Party are each a “Party” and together "
    f"the “Parties.”"
)

SECTIONS = [
    ("1. Purpose",
     "The Parties wish to explore a potential business relationship (the “Purpose”). "
     "In connection with the Purpose, the Disclosing Party may disclose to the Receiving Party "
     "certain confidential and proprietary information."),
    ("2. Confidential Information",
     "“Confidential Information” means any non-public information disclosed by the "
     "Disclosing Party to the Receiving Party, whether orally, in writing, or in any other form, "
     "that is designated as confidential or that reasonably should be understood to be confidential "
     "given its nature and the circumstances of disclosure. It includes, without limitation, business "
     "plans, financials, pricing, technology, software, source code, models, prompts, algorithms, "
     "product and roadmap information, clinical and operational data, customer and patient information, "
     "trade secrets, and know-how."),
    ("3. Obligations of the Receiving Party",
     "The Receiving Party shall: (a) hold the Confidential Information in strict confidence; "
     "(b) not disclose it to any third party without the Disclosing Party’s prior written consent; "
     "(c) use it solely for the Purpose; (d) protect it using at least the same degree of care it uses "
     "for its own confidential information, and in no event less than reasonable care; and (e) limit "
     "access to its personnel and advisors who need to know it for the Purpose and who are bound by "
     "confidentiality obligations at least as protective as those herein."),
    ("4. Exclusions",
     "Confidential Information does not include information that: (a) is or becomes publicly available "
     "through no fault of the Receiving Party; (b) was rightfully known to the Receiving Party without "
     "restriction before disclosure; (c) is rightfully received from a third party without a duty of "
     "confidentiality; or (d) is independently developed by the Receiving Party without use of or "
     "reference to the Confidential Information. If disclosure is required by law or court order, the "
     "Receiving Party may disclose only as required, after giving the Disclosing Party prompt notice "
     "(where legally permitted) and reasonable cooperation to seek protective treatment."),
    ("5. Term and Survival",
     "This Agreement applies to Confidential Information disclosed during its term, which begins on the "
     "Effective Date and continues until terminated by either Party on written notice. The Receiving "
     "Party’s confidentiality obligations survive for [3] years after disclosure, except that "
     "obligations with respect to trade secrets continue for as long as the information remains a trade "
     "secret under applicable law."),
    ("6. Return or Destruction",
     "Upon the Disclosing Party’s written request, the Receiving Party shall promptly return or "
     "destroy all Confidential Information and any copies, and, if requested, certify such destruction "
     "in writing."),
    ("7. No License or Rights",
     "No license or other right to any intellectual property is granted under this Agreement. All "
     "Confidential Information remains the property of the Disclosing Party."),
    ("8. No Warranty",
     "All Confidential Information is provided “as is.” The Disclosing Party makes no "
     "warranties, express or implied, regarding its accuracy or completeness."),
    ("9. No Obligation",
     "Nothing in this Agreement obligates either Party to proceed with any transaction or relationship."),
    ("10. Remedies",
     "The Receiving Party agrees that any breach may cause irreparable harm for which monetary damages "
     "would be inadequate, and that the Disclosing Party is entitled to seek injunctive relief in "
     "addition to any other remedies available at law or in equity."),
    ("11. Governing Law",
     "This Agreement is governed by the laws of [State/Jurisdiction], without regard to its conflict-of-"
     "laws principles."),
    ("12. Miscellaneous",
     "This Agreement is the entire agreement between the Parties regarding its subject matter and "
     "supersedes all prior understandings. It may be amended only in a writing signed by both Parties. "
     "If any provision is held unenforceable, the remaining provisions remain in effect. The Receiving "
     "Party may not assign this Agreement without the Disclosing Party’s prior written consent."),
]

SIGN_INTRO = "IN WITNESS WHEREOF, the Receiving Party has executed this Agreement as of the date below."


def get_secret():
    pw = os.environ.get("NDA_SECRET")
    if pw:
        return pw
    envp = os.path.join(REPO, ".env")
    if os.path.exists(envp):
        for line in open(envp, encoding="utf-8"):
            line = line.strip()
            if line.startswith("NDA_SECRET="):
                return line.split("=", 1)[1].strip().strip('"').strip("'")
    return None


def pdf_safe(s):
    """fpdf2's core Helvetica is latin-1 only; map smart punctuation to ASCII."""
    return (s.replace("“", '"').replace("”", '"')
             .replace("‘", "'").replace("’", "'")
             .replace("—", "-").replace("–", "-"))


def build_pdf(password):
    pdf = FPDF(format="letter")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(20, 18, 20)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 9, TITLE, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, SUBTITLE, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 10.5)
    pdf.multi_cell(0, 5.4, pdf_safe(INTRO), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    for heading, body in SECTIONS:
        pdf.set_font("Helvetica", "B", 11)
        pdf.multi_cell(0, 6, pdf_safe(heading), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10.5)
        pdf.multi_cell(0, 5.4, pdf_safe(body), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1.5)

    pdf.ln(3)
    pdf.set_font("Helvetica", "", 10.5)
    pdf.multi_cell(0, 5.4, pdf_safe(SIGN_INTRO), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)
    for label in ("Receiving Party (print name): ______________________________",
                  "Signature: ______________________________",
                  "Title / Organization: ______________________________",
                  "Date: ______________________________"):
        pdf.cell(0, 8, label, new_x="LMARGIN", new_y="NEXT")

    # Open password = NDA_SECRET (recipients must enter it to view the document).
    pdf.set_encryption(owner_password=password, user_password=password,
                       permissions=AccessPermission.all())
    pdf.output(PDF_OUT)


def build_docx(password):
    doc = Document()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(SUBTITLE).font.size = Pt(10)

    doc.add_paragraph(INTRO)
    for heading, body in SECTIONS:
        hp = doc.add_paragraph()
        hp.add_run(heading).bold = True
        doc.add_paragraph(body)

    doc.add_paragraph()
    doc.add_paragraph(SIGN_INTRO)
    for label in ("Receiving Party (print name): ______________________________",
                  "Signature: ______________________________",
                  "Title / Organization: ______________________________",
                  "Date: ______________________________"):
        doc.add_paragraph(label)

    # Save to memory, then encrypt with the open password = NDA_SECRET.
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with open(DOCX_OUT, "wb") as out:
        OOXMLFile(buf).encrypt(password, out)


def main():
    pw = get_secret()
    if not pw:
        sys.exit("NDA_SECRET is not set. Add NDA_SECRET=<password> to the repo .env "
                 "(or export it), then re-run.")
    build_pdf(pw)
    build_docx(pw)
    print(f"Wrote (encrypted with NDA_SECRET):\n  {PDF_OUT}\n  {DOCX_OUT}")


if __name__ == "__main__":
    main()
