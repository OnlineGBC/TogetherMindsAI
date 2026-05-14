"""
Convert README.md to README.docx using python-docx.

Run from repo root:
    python scripts/readme_to_docx.py
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


REPO_ROOT = Path(__file__).resolve().parent.parent
SRC = REPO_ROOT / "README.md"
DST = REPO_ROOT / "README.docx"


def add_inline(paragraph, text: str) -> None:
    """Render a single line of inline markdown (bold, code, links) into a paragraph."""
    # Pattern matches: **bold**, `code`, [text](url)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                run = paragraph.add_run(link_match.group(1))
                run.font.color.rgb = RGBColor(0x06, 0x5F, 0xD2)
                run.underline = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_to_docx(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Light Grid Accent 1"
        for r, row in enumerate(table_rows):
            for c, cell_text in enumerate(row):
                cell = table.rows[r].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline(p, cell_text)
                if r == 0:
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()
        in_table = False
        table_rows = []

    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Tables
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Skip separator rows like |---|---|
            if all(re.fullmatch(r":?-+:?", c) for c in cells):
                i += 1
                continue
            in_table = True
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, line[2:].strip())
        elif re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, text)
        elif re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, text)
        elif line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1

    if in_table:
        flush_table()


def main() -> None:
    md_text = SRC.read_text(encoding="utf-8")
    doc = Document()

    # Tighten default style spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    md_to_docx(md_text, doc)
    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
